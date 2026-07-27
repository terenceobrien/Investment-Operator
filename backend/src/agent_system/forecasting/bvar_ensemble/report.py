"""Word report generation for saved BVAR ensemble forecasts."""
from __future__ import annotations

import json
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd

from src.agent_system.forecasting.bvar_ensemble.estimation import default_bvar_cache_dir
from src.agent_system.forecasting.scenario_classifier.registry import VariableRegistry
from src.agent_system.paths import project_root


class ReportError(RuntimeError):
    """Raised when a saved forecast cannot be rendered as a report."""


SCENARIO_PALETTE = {
    "expansion_disinflation": "#4C78A8",
    "late_cycle_expansion": "#72B7B2",
    "inflation_shock": "#F58518",
    "stagflation": "#B279A2",
    "growth_scare_no_credit": "#54A24B",
    "credit_led_recession": "#E45756",
}

TERMINAL_HISTOGRAM_VARIABLES = ["activity", "lur", "core_pce", "credit_spread"]
TRANSFORM_LABELS = {
    "level": "level",
    "yoy_pct": "yoy %",
    "qoq_ann_pct": "q/q ann %",
    "diff": "diff",
}


def newest_forecast_json(*, bvar_cache_dir: str | Path | None = None) -> Path:
    cache_dir = Path(bvar_cache_dir) if bvar_cache_dir is not None else default_bvar_cache_dir()
    if not cache_dir.is_dir():
        raise ReportError(f"BVAR cache directory not found: {cache_dir}")
    candidates = sorted(
        cache_dir.glob("forecast_*.json"),
        key=lambda path: (path.stat().st_mtime, path.name),
        reverse=True,
    )
    if not candidates:
        raise ReportError(f"No forecast_*.json files found in {cache_dir}")
    return candidates[0]


def resolve_report_output_dir(
    config: dict[str, Any],
    *,
    bvar_cache_dir: str | Path | None = None,
) -> Path:
    raw = str(config.get("report_output_dir") or "").strip()
    default_literal = "data/agent_system/bvar_cache/reports"
    if bvar_cache_dir is not None and (not raw or raw == default_literal):
        return Path(bvar_cache_dir) / "reports"
    if not raw:
        return default_bvar_cache_dir() / "reports"
    path = Path(raw)
    if path.is_absolute():
        return path
    return project_root() / path


def generate_forecast_report(
    forecast_path: str | Path | None = None,
    *,
    compare_forecast_path: str | Path | None = None,
    output_dir: str | Path | None = None,
    bvar_cache_dir: str | Path | None = None,
) -> Path:
    forecast_file = Path(forecast_path) if forecast_path is not None else newest_forecast_json(
        bvar_cache_dir=bvar_cache_dir
    )
    if not forecast_file.is_file():
        raise ReportError(f"forecast JSON not found: {forecast_file}")
    forecast = _read_json(forecast_file)
    if compare_forecast_path is not None:
        return _generate_comparison_report(
            forecast_file,
            forecast,
            Path(compare_forecast_path),
            output_dir=output_dir,
        )
    paths_path = _simulation_paths_path(forecast, forecast_file)
    paths = _read_paths(paths_path)
    registry = VariableRegistry.load()
    variable_specs = {spec.name: spec for spec in registry.spine_variables()}
    spine_variables = registry.spine_variable_names()
    missing = [variable for variable in spine_variables if variable not in paths.columns]
    if missing:
        raise ReportError(
            f"simulation paths parquet missing spine variables {missing}: {paths_path}"
        )

    docx = _docx_imports()
    Document = docx["Document"]
    Inches = docx["Inches"]
    Pt = docx["Pt"]

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    doc.add_heading("BVAR Ensemble Forecast", level=0)
    _add_header_table(doc, forecast)
    _add_shock_note(doc, forecast)
    _add_model_limitations(doc, forecast)
    _add_scenario_section(doc, forecast, Inches)

    with tempfile.TemporaryDirectory(prefix="bvar_report_") as tmp:
        tmp_dir = Path(tmp)
        chart_path = tmp_dir / "scenario_probabilities.png"
        _plot_scenario_probabilities(forecast, chart_path)
        doc.add_picture(str(chart_path), width=Inches(6.4))
        _add_margin_notes(doc, forecast)
        _add_regime_summary(doc, forecast)

        doc.add_heading("Fan Charts", level=1)
        for variable in spine_variables:
            spec = variable_specs[variable]
            chart = tmp_dir / f"fan_{variable}.png"
            _plot_fan_chart(
                paths,
                forecast=forecast,
                variable=variable,
                transform=spec.transform,
                chart_path=chart,
            )
            doc.add_picture(str(chart), width=Inches(6.4))

        doc.add_heading("Terminal Distributions", level=1)
        for variable in TERMINAL_HISTOGRAM_VARIABLES:
            if variable not in paths.columns:
                continue
            spec = registry.get(variable)
            chart = tmp_dir / f"hist_{variable}.png"
            _plot_terminal_histogram(
                paths,
                forecast=forecast,
                variable=variable,
                transform=spec.transform,
                chart_path=chart,
            )
            doc.add_picture(str(chart), width=Inches(6.4))

    _add_tail_diagnostics(doc, forecast)
    _add_validity_summary(doc, forecast)
    _add_appendix(doc, forecast)

    props = doc.core_properties
    props.author = "AI Financial Operator"
    props.title = "BVAR Ensemble Forecast"
    fixed_time = _forecast_datetime(forecast)
    props.created = fixed_time
    props.modified = fixed_time

    target_dir = Path(output_dir) if output_dir is not None else default_bvar_cache_dir() / "reports"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / (
        f"bvar_forecast_{forecast.get('asof_quarter', 'unknown')}_"
        f"{_safe_timestamp(str(forecast.get('generated_at') or _utc_now()))}.docx"
    )
    doc.save(out_path)
    return out_path


def _generate_comparison_report(
    primary_file: Path,
    primary: dict[str, Any],
    compare_file: Path,
    *,
    output_dir: str | Path | None,
) -> Path:
    if not compare_file.is_file():
        raise ReportError(f"comparison forecast JSON not found: {compare_file}")
    comparison = _read_json(compare_file)
    primary_paths = _read_paths(_simulation_paths_path(primary, primary_file))
    comparison_paths = _read_paths(_simulation_paths_path(comparison, compare_file))
    registry = VariableRegistry.load()
    spine_variables = registry.spine_variable_names()

    docx = _docx_imports()
    Document = docx["Document"]
    Inches = docx["Inches"]
    Pt = docx["Pt"]
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    doc.add_heading("BVAR Ensemble Forecast Comparison", level=0)
    primary_label = _forecast_label(primary, "Primary")
    comparison_label = _forecast_label(comparison, "Comparison")
    doc.add_paragraph(
        f"Primary: {primary_label}. Comparison: {comparison_label}."
    )
    _add_comparison_header_table(doc, primary, comparison)
    _add_model_limitations(doc, primary)
    _add_scenario_comparison_section(doc, primary, comparison)
    _add_regime_comparison_summary(doc, primary, comparison)

    with tempfile.TemporaryDirectory(prefix="bvar_report_compare_") as tmp:
        tmp_dir = Path(tmp)
        chart_path = tmp_dir / "scenario_probability_compare.png"
        _plot_scenario_probability_comparison(primary, comparison, chart_path)
        doc.add_picture(str(chart_path), width=Inches(6.4))

        doc.add_heading("Fan Chart Comparison", level=1)
        for variable in spine_variables:
            spec = registry.get(variable)
            chart = tmp_dir / f"fan_compare_{variable}.png"
            _plot_fan_comparison(
                primary_paths,
                comparison_paths,
                primary=primary,
                comparison=comparison,
                variable=variable,
                transform=spec.transform,
                chart_path=chart,
            )
            doc.add_picture(str(chart), width=Inches(6.4))

        doc.add_heading("Terminal Distribution Comparison", level=1)
        for variable in TERMINAL_HISTOGRAM_VARIABLES:
            spec = registry.get(variable)
            chart = tmp_dir / f"hist_compare_{variable}.png"
            _plot_terminal_histogram_comparison(
                primary_paths,
                comparison_paths,
                primary=primary,
                comparison=comparison,
                variable=variable,
                transform=spec.transform,
                chart_path=chart,
            )
            doc.add_picture(str(chart), width=Inches(6.4))

    _add_appendix(doc, primary)
    doc.core_properties.author = "AI Financial Operator"
    doc.core_properties.title = "BVAR Ensemble Forecast Comparison"
    fixed_time = _forecast_datetime(primary)
    doc.core_properties.created = fixed_time
    doc.core_properties.modified = fixed_time
    target_dir = Path(output_dir) if output_dir is not None else default_bvar_cache_dir() / "reports"
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / (
        f"bvar_forecast_compare_{primary.get('asof_quarter', 'unknown')}_"
        f"{_safe_timestamp(str(primary.get('generated_at') or _utc_now()))}.docx"
    )
    doc.save(out_path)
    return out_path


def _add_header_table(doc: Any, forecast: dict[str, Any]) -> None:
    rows = [
        ("As-of quarter", forecast.get("asof_quarter")),
        ("Generated at", forecast.get("generated_at")),
        ("Paths", forecast.get("n_paths")),
        ("Horizon", forecast.get("horizon_quarters")),
        ("Shock distribution", forecast.get("shock_dist")),
        ("Vol model", forecast.get("vol_model")),
        ("Regime model", forecast.get("regime_model")),
        ("Anchor regime p_enter", forecast.get("regime_anchor_p_enter")),
        ("Regime artifact", forecast.get("regime_artifact")),
        ("Seed", forecast.get("seed")),
        ("Baseline mode", forecast.get("baseline_mode")),
        ("Posterior artifact", forecast.get("posterior_artifact")),
        ("Posterior fingerprint", forecast.get("posterior_artifact_fingerprint")),
        ("Handoff fingerprint", forecast.get("handoff_fingerprint")),
        ("Handoff file", forecast.get("handoff_file")),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if value is None else str(value)


def _add_comparison_header_table(
    doc: Any,
    primary: dict[str, Any],
    comparison: dict[str, Any],
) -> None:
    rows = [
        ("As-of quarter", primary.get("asof_quarter"), comparison.get("asof_quarter")),
        ("Generated at", primary.get("generated_at"), comparison.get("generated_at")),
        ("Paths", primary.get("n_paths"), comparison.get("n_paths")),
        ("Horizon", primary.get("horizon_quarters"), comparison.get("horizon_quarters")),
        ("Vol model", primary.get("vol_model"), comparison.get("vol_model")),
        ("Regime model", primary.get("regime_model"), comparison.get("regime_model")),
        ("Anchor regime p_enter", primary.get("regime_anchor_p_enter"), comparison.get("regime_anchor_p_enter")),
        ("Shock distribution", primary.get("shock_dist"), comparison.get("shock_dist")),
        ("Seed", primary.get("seed"), comparison.get("seed")),
        ("Baseline mode", primary.get("baseline_mode"), comparison.get("baseline_mode")),
        ("Posterior fingerprint", primary.get("posterior_artifact_fingerprint"), comparison.get("posterior_artifact_fingerprint")),
        ("Handoff fingerprint", primary.get("handoff_fingerprint"), comparison.get("handoff_fingerprint")),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Primary"
    table.rows[0].cells[2].text = "Comparison"
    for key, left, right in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if left is None else str(left)
        cells[2].text = "" if right is None else str(right)


def _add_shock_note(doc: Any, forecast: dict[str, Any]) -> None:
    shock_dist = str(forecast.get("shock_dist") or "")
    if shock_dist == "student_t":
        doc.add_paragraph(
            "Note: this run used Student-t shocks, so the ensemble intentionally carries fatter simulated tails than the Gaussian setting."
        )
    elif shock_dist == "gaussian":
        doc.add_paragraph(
            "Note: this run used Gaussian shocks."
        )


def _add_model_limitations(doc: Any, forecast: dict[str, Any]) -> None:
    limitations = forecast.get("model_limitations") or {}
    if not isinstance(limitations, dict) or not limitations:
        return
    magnitude = limitations.get("credit_tail_magnitude")
    detail = limitations.get("detail")
    doc.add_heading("Model Limitations", level=1)
    if magnitude:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Credit tail magnitude: ")
        run.bold = True
        paragraph.add_run(str(magnitude))
    if detail:
        doc.add_paragraph(str(detail))


def _add_scenario_section(doc: Any, forecast: dict[str, Any], Inches: Any) -> None:
    del Inches
    doc.add_heading("Scenario Probabilities", level=1)
    hard = forecast.get("scenario_probabilities") or {}
    soft = forecast.get("scenario_probabilities_soft") or {}
    rows = sorted(hard.items(), key=lambda item: float(item[1]), reverse=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    headers = table.rows[0].cells
    headers[0].text = "Scenario"
    headers[1].text = "Hard primary"
    headers[2].text = "Soft secondary"
    for scenario, probability in rows:
        cells = table.add_row().cells
        cells[0].text = scenario
        cells[1].text = f"{float(probability):.1%}"
        cells[2].text = f"{float(soft.get(scenario, 0.0)):.1%}"


def _add_scenario_comparison_section(
    doc: Any,
    primary: dict[str, Any],
    comparison: dict[str, Any],
) -> None:
    doc.add_heading("Scenario Probabilities", level=1)
    primary_probs = primary.get("scenario_probabilities") or {}
    comparison_probs = comparison.get("scenario_probabilities") or {}
    scenarios = sorted(
        set(primary_probs) | set(comparison_probs),
        key=lambda scenario: float(primary_probs.get(scenario, 0.0)),
        reverse=True,
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    headers = table.rows[0].cells
    headers[0].text = "Scenario"
    headers[1].text = "Primary"
    headers[2].text = "Comparison"
    headers[3].text = "Delta"
    for scenario in scenarios:
        left = float(primary_probs.get(scenario, 0.0))
        right = float(comparison_probs.get(scenario, 0.0))
        cells = table.add_row().cells
        cells[0].text = scenario
        cells[1].text = f"{left:.1%}"
        cells[2].text = f"{right:.1%}"
        cells[3].text = f"{left - right:+.1%}"


def _add_margin_notes(doc: Any, forecast: dict[str, Any]) -> None:
    stats = forecast.get("margin_stats") or {}
    doc.add_paragraph(
        "Margin mean: "
        f"{float(stats.get('mean', 0.0)):.3f}; "
        f"p25: {float(stats.get('p25', 0.0)):.3f}; "
        f"share low margin: {float(stats.get('share_low_margin', 0.0)):.1%}."
    )
    doc.add_paragraph(
        "High share_low_margin means many simulated futures were hard for the classifier to separate cleanly."
    )


def _add_regime_summary(doc: Any, forecast: dict[str, Any]) -> None:
    if str(forecast.get("regime_model") or "none") != "markov":
        return
    doc.add_heading("Regime Overlay", level=1)
    diagnostics = forecast.get("regime_diagnostics") or {}
    rows = [
        ("Anchor label", forecast.get("regime_anchor_label")),
        ("Anchor p_enter", _fmt_probability(forecast.get("regime_anchor_p_enter"))),
        ("Fraction entered stress", _fmt_probability(diagnostics.get("fraction_entered_stress"))),
        ("Fraction ever stress", _fmt_probability(diagnostics.get("fraction_ever_stress"))),
        ("Avg quarters in stress", _fmt_float(diagnostics.get("avg_quarters_in_stress"))),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if value is None else str(value)
    split = diagnostics.get("credit_spread_terminal_by_regime_path") or {}
    if split:
        doc.add_paragraph("Credit spread terminal distribution split by regime path:")
        split_table = doc.add_table(rows=1, cols=6)
        split_table.style = "Light Grid"
        for cell, header in zip(
            split_table.rows[0].cells,
            ["Group", "Count", "p50", "p90", "p99", "Max"],
        ):
            cell.text = header
        for group in ["ever_stress", "stayed_calm"]:
            payload = split.get(group) or {}
            cells = split_table.add_row().cells
            cells[0].text = group
            cells[1].text = str(payload.get("count", 0))
            cells[2].text = _fmt_float(payload.get("p50"))
            cells[3].text = _fmt_float(payload.get("p90"))
            cells[4].text = _fmt_float(payload.get("p99"))
            cells[5].text = _fmt_float(payload.get("max"))


def _add_regime_comparison_summary(
    doc: Any,
    primary: dict[str, Any],
    comparison: dict[str, Any],
) -> None:
    if (
        str(primary.get("regime_model") or "none") == "none"
        and str(comparison.get("regime_model") or "none") == "none"
    ):
        return
    doc.add_heading("Regime Overlay Comparison", level=1)
    metrics = [
        ("Regime model", primary.get("regime_model"), comparison.get("regime_model")),
        ("Anchor label", primary.get("regime_anchor_label"), comparison.get("regime_anchor_label")),
        ("Anchor p_enter", _fmt_probability(primary.get("regime_anchor_p_enter")), _fmt_probability(comparison.get("regime_anchor_p_enter"))),
        ("Fraction entered stress", _fmt_probability((primary.get("regime_diagnostics") or {}).get("fraction_entered_stress")), _fmt_probability((comparison.get("regime_diagnostics") or {}).get("fraction_entered_stress"))),
        ("Avg quarters in stress", _fmt_float((primary.get("regime_diagnostics") or {}).get("avg_quarters_in_stress")), _fmt_float((comparison.get("regime_diagnostics") or {}).get("avg_quarters_in_stress"))),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Primary"
    table.rows[0].cells[2].text = "Comparison"
    for key, left, right in metrics:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if left is None else str(left)
        cells[2].text = "" if right is None else str(right)


def _add_tail_diagnostics(doc: Any, forecast: dict[str, Any]) -> None:
    doc.add_heading("Tail Diagnostics", level=1)
    doc.add_paragraph(
        "Ensemble K-quarter changes are compared with historical K-quarter changes; these diagnostics flag unusual tails but do not filter paths."
    )
    diagnostics = forecast.get("tail_diagnostics") or {}
    table = doc.add_table(rows=1, cols=9)
    table.style = "Light Grid"
    headers = [
        "Variable",
        "Ens p50",
        "Ens p90",
        "Ens p99",
        "Ens max",
        "Hist p50",
        "Hist p90",
        "Hist p99",
        "Hist max",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for variable, payload in diagnostics.items():
        row = table.add_row()
        ensemble = payload.get("ensemble") or {}
        historical = payload.get("historical") or {}
        values = [
            variable,
            ensemble.get("p50"),
            ensemble.get("p90"),
            ensemble.get("p99"),
            ensemble.get("max"),
            historical.get("p50"),
            historical.get("p90"),
            historical.get("p99"),
            historical.get("max"),
        ]
        for cell, value in zip(row.cells, values):
            cell.text = str(value) if isinstance(value, str) else _fmt_float(value)
        if payload.get("flag_ensemble_p99_exceeds_historical_max"):
            _bold_row(row)
            _shade_row(row, "FCE4D6")


def _add_validity_summary(doc: Any, forecast: dict[str, Any]) -> None:
    doc.add_heading("Validity Summary", level=1)
    validity = forecast.get("validity") or {}
    rows = [
        ("Rejections", validity.get("rejections")),
        ("Redraws", validity.get("redraws")),
        ("Clips", validity.get("clips")),
        ("Rejection rate", f"{float(validity.get('rejection_rate_pct', 0.0)):.2f}%"),
        ("Warning threshold", f"{float(validity.get('rejection_warn_pct', 0.0)):.2f}%"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = "" if value is None else str(value)
    warning = validity.get("warning")
    if warning:
        doc.add_paragraph(str(warning))
    violations = validity.get("per_variable_violations") or {}
    doc.add_paragraph("Per-variable violation counts:")
    vtable = doc.add_table(rows=1, cols=2)
    vtable.style = "Light Grid"
    vtable.rows[0].cells[0].text = "Variable"
    vtable.rows[0].cells[1].text = "Violations"
    if violations:
        for variable, count in sorted(violations.items()):
            cells = vtable.add_row().cells
            cells[0].text = str(variable)
            cells[1].text = str(count)
    else:
        cells = vtable.add_row().cells
        cells[0].text = "none"
        cells[1].text = "0"


def _add_appendix(doc: Any, forecast: dict[str, Any]) -> None:
    doc.add_heading("Appendix: Run Configuration", level=1)
    rows: list[tuple[str, Any]] = []
    config = forecast.get("config") or {}
    for key, value in sorted(config.items()):
        rows.append((f"config.{key}", value))
    for key, value in sorted((forecast.get("posterior_hyperparameters") or {}).items()):
        rows.append((f"posterior.{key}", value))
    classifier_metadata = forecast.get("classifier_metadata") or {}
    for key in ["scaling", "kernel_sigma", "active_variables", "handoff_file"]:
        if key in classifier_metadata:
            rows.append((f"classifier.{key}", classifier_metadata[key]))
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = json.dumps(value, sort_keys=True) if isinstance(value, (dict, list)) else str(value)


def _plot_scenario_probabilities(forecast: dict[str, Any], chart_path: Path) -> None:
    plt = _matplotlib_pyplot()
    hard = forecast.get("scenario_probabilities") or {}
    rows = sorted(hard.items(), key=lambda item: float(item[1]))
    scenarios = [scenario for scenario, _ in rows]
    values = [float(value) for _, value in rows]
    colors = [SCENARIO_PALETTE.get(scenario, "#6B7280") for scenario in scenarios]
    fig, ax = plt.subplots(figsize=(7.0, 3.8))
    ax.barh(scenarios, values, color=colors)
    ax.set_xlim(0, max(0.01, max(values) * 1.15 if values else 1.0))
    ax.set_xlabel("Probability")
    ax.grid(axis="x", alpha=0.22)
    ax.set_title("Hard Scenario Probabilities")
    for index, value in enumerate(values):
        ax.text(value + 0.005, index, f"{value:.1%}", va="center", fontsize=8)
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _plot_scenario_probability_comparison(
    primary: dict[str, Any],
    comparison: dict[str, Any],
    chart_path: Path,
) -> None:
    plt = _matplotlib_pyplot()
    primary_probs = primary.get("scenario_probabilities") or {}
    comparison_probs = comparison.get("scenario_probabilities") or {}
    scenarios = sorted(
        set(primary_probs) | set(comparison_probs),
        key=lambda scenario: float(primary_probs.get(scenario, 0.0)),
    )
    y = np.arange(len(scenarios))
    left = np.asarray([float(primary_probs.get(scenario, 0.0)) for scenario in scenarios])
    right = np.asarray([float(comparison_probs.get(scenario, 0.0)) for scenario in scenarios])
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    ax.barh(y + 0.18, left, height=0.32, color="#4C78A8", label=_forecast_label(primary, "Primary"))
    ax.barh(y - 0.18, right, height=0.32, color="#F58518", label=_forecast_label(comparison, "Comparison"))
    ax.set_yticks(y)
    ax.set_yticklabels(scenarios)
    ax.set_xlabel("Probability")
    ax.grid(axis="x", alpha=0.22)
    ax.legend(frameon=False, fontsize=8)
    ax.set_title("Hard Scenario Probabilities: Primary vs Comparison")
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _plot_fan_chart(
    paths: pd.DataFrame,
    *,
    forecast: dict[str, Any],
    variable: str,
    transform: str,
    chart_path: Path,
) -> None:
    plt = _matplotlib_pyplot()
    matrix = _variable_matrix(paths, variable, int(forecast["horizon_quarters"]))
    anchor = float((forecast.get("anchor_values") or {})[variable])
    quantiles = np.percentile(matrix, [10, 25, 50, 75, 90], axis=0)
    q10, q25, q50, q75, q90 = [
        np.concatenate([[anchor], values])
        for values in quantiles
    ]
    x = np.arange(0, matrix.shape[1] + 1)
    labels = _quarter_labels(str(forecast["asof_quarter"]), matrix.shape[1])
    fig, ax = plt.subplots(figsize=(7.0, 3.8))
    ax.fill_between(x, q10, q90, color="#9CA3AF", alpha=0.28, label="p10-p90")
    ax.fill_between(x, q25, q75, color="#4C78A8", alpha=0.25, label="p25-p75")
    ax.plot(x, q50, color="#1F2937", lw=2.0, label="median")
    ax.scatter([0], [anchor], color="#111827", s=20, zorder=4)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=0, fontsize=8)
    ax.set_title(f"{variable} ({TRANSFORM_LABELS.get(transform, transform)})")
    ax.grid(alpha=0.22)
    ax.legend(loc="best", fontsize=8, frameon=False)
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _plot_fan_comparison(
    primary_paths: pd.DataFrame,
    comparison_paths: pd.DataFrame,
    *,
    primary: dict[str, Any],
    comparison: dict[str, Any],
    variable: str,
    transform: str,
    chart_path: Path,
) -> None:
    plt = _matplotlib_pyplot()
    horizon = int(primary["horizon_quarters"])
    left_matrix = _variable_matrix(primary_paths, variable, horizon)
    right_matrix = _variable_matrix(comparison_paths, variable, int(comparison["horizon_quarters"]))
    left_anchor = float((primary.get("anchor_values") or {})[variable])
    right_anchor = float((comparison.get("anchor_values") or {})[variable])
    x = np.arange(0, horizon + 1)
    labels = _quarter_labels(str(primary["asof_quarter"]), horizon)
    left_q = _fan_quantiles_with_anchor(left_matrix, left_anchor)
    right_q = _fan_quantiles_with_anchor(right_matrix, right_anchor)
    fig, ax = plt.subplots(figsize=(7.0, 3.8))
    ax.fill_between(x, left_q["p10"], left_q["p90"], color="#4C78A8", alpha=0.18)
    ax.plot(x, left_q["p50"], color="#1F2937", lw=2.0, label=_forecast_label(primary, "Primary"))
    ax.fill_between(x, right_q["p10"], right_q["p90"], color="#F58518", alpha=0.16)
    ax.plot(x, right_q["p50"], color="#B45309", lw=1.8, label=_forecast_label(comparison, "Comparison"))
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_title(f"{variable} fan comparison ({TRANSFORM_LABELS.get(transform, transform)})")
    ax.grid(alpha=0.22)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _plot_terminal_histogram(
    paths: pd.DataFrame,
    *,
    forecast: dict[str, Any],
    variable: str,
    transform: str,
    chart_path: Path,
) -> None:
    plt = _matplotlib_pyplot()
    horizon = int(forecast["horizon_quarters"])
    terminal = paths.loc[paths["quarter_index"] == horizon, variable].to_numpy(dtype=float)
    anchor = float((forecast.get("anchor_values") or {})[variable])
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.hist(terminal, bins=40, color="#4C78A8", alpha=0.78)
    ax.axvline(anchor, color="#E45756", lw=2.0, label="anchor")
    ax.set_title(
        f"{variable} terminal distribution ({TRANSFORM_LABELS.get(transform, transform)})"
    )
    ax.grid(axis="y", alpha=0.22)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _plot_terminal_histogram_comparison(
    primary_paths: pd.DataFrame,
    comparison_paths: pd.DataFrame,
    *,
    primary: dict[str, Any],
    comparison: dict[str, Any],
    variable: str,
    transform: str,
    chart_path: Path,
) -> None:
    plt = _matplotlib_pyplot()
    left_horizon = int(primary["horizon_quarters"])
    right_horizon = int(comparison["horizon_quarters"])
    left = primary_paths.loc[
        primary_paths["quarter_index"] == left_horizon,
        variable,
    ].to_numpy(dtype=float)
    right = comparison_paths.loc[
        comparison_paths["quarter_index"] == right_horizon,
        variable,
    ].to_numpy(dtype=float)
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    bins = 40
    ax.hist(left, bins=bins, color="#4C78A8", alpha=0.42, density=True, label=_forecast_label(primary, "Primary"))
    ax.hist(right, bins=bins, color="#F58518", alpha=0.38, density=True, label=_forecast_label(comparison, "Comparison"))
    ax.axvline(float((primary.get("anchor_values") or {})[variable]), color="#1F2937", lw=1.5, label="primary anchor")
    ax.set_title(
        f"{variable} terminal distribution comparison ({TRANSFORM_LABELS.get(transform, transform)})"
    )
    ax.grid(axis="y", alpha=0.22)
    ax.legend(frameon=False, fontsize=8)
    fig.tight_layout()
    fig.savefig(chart_path, dpi=160)
    plt.close(fig)


def _fan_quantiles_with_anchor(matrix: np.ndarray, anchor: float) -> dict[str, np.ndarray]:
    quantiles = np.percentile(matrix, [10, 50, 90], axis=0)
    return {
        "p10": np.concatenate([[anchor], quantiles[0]]),
        "p50": np.concatenate([[anchor], quantiles[1]]),
        "p90": np.concatenate([[anchor], quantiles[2]]),
    }


def _forecast_label(forecast: dict[str, Any], fallback: str) -> str:
    vol_model = forecast.get("vol_model") or "constant"
    regime_model = forecast.get("regime_model") or "none"
    shock = forecast.get("shock_dist") or "unknown"
    seed = forecast.get("seed")
    return f"{vol_model}/{regime_model}/{shock}/seed {seed}" if seed is not None else fallback


def _variable_matrix(paths: pd.DataFrame, variable: str, horizon: int) -> np.ndarray:
    required = {"path_id", "quarter_index", variable}
    missing = sorted(required - set(paths.columns))
    if missing:
        raise ReportError(f"paths parquet missing required columns: {missing}")
    pivot = paths.pivot(index="path_id", columns="quarter_index", values=variable)
    expected_columns = list(range(1, horizon + 1))
    missing_quarters = [quarter for quarter in expected_columns if quarter not in pivot.columns]
    if missing_quarters:
        raise ReportError(f"paths parquet missing quarters for {variable}: {missing_quarters}")
    matrix = pivot[expected_columns].sort_index().to_numpy(dtype=float)
    if matrix.size == 0 or not np.isfinite(matrix).all():
        raise ReportError(f"paths parquet contains non-finite values for {variable}")
    return matrix


def _read_json(path: Path) -> dict[str, Any]:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise ReportError(f"could not read forecast JSON {path}: {exc}") from exc
    if not isinstance(payload, dict):
        raise ReportError(f"forecast JSON must contain an object: {path}")
    return payload


def _simulation_paths_path(forecast: dict[str, Any], forecast_file: Path) -> Path:
    raw = forecast.get("simulation_paths_parquet")
    if raw:
        path = Path(str(raw))
        if not path.is_absolute():
            if path.is_file():
                path = path
            else:
                path = forecast_file.parent / path
    else:
        path = forecast_file.with_name(f"{forecast_file.stem}_spine_paths.parquet")
    if not path.is_file():
        raise ReportError(
            "simulation paths parquet missing for report. "
            "Run a new BVAR forecast so the JSON includes simulation_paths_parquet."
        )
    return path


def _read_paths(path: Path) -> pd.DataFrame:
    try:
        frame = pd.read_parquet(path)
    except Exception as exc:
        raise ReportError(f"could not read paths parquet {path}: {exc}") from exc
    required = {"path_id", "quarter_index"}
    missing = sorted(required - set(frame.columns))
    if missing:
        raise ReportError(f"paths parquet missing required columns {missing}: {path}")
    return frame


def _quarter_labels(asof_quarter: str, horizon: int) -> list[str]:
    anchor = pd.Period(asof_quarter, freq="Q")
    return [str(anchor + offset) for offset in range(0, horizon + 1)]


def _docx_imports() -> dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:
        raise ReportError("python-docx is required for report generation") from exc
    return {"Document": Document, "Inches": Inches, "Pt": Pt}


def _matplotlib_pyplot() -> Any:
    mpl_dir = Path(tempfile.gettempdir()) / "ai_financial_operator_matplotlib"
    mpl_dir.mkdir(parents=True, exist_ok=True)
    os.environ.setdefault("MPLCONFIGDIR", str(mpl_dir))
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    plt.rcParams.update(
        {
            "axes.spines.top": False,
            "axes.spines.right": False,
            "font.size": 9,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )
    return plt


def _bold_row(row: Any) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _shade_row(row: Any, fill: str) -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    for cell in row.cells:
        cell._tc.get_or_add_tcPr().append(  # noqa: SLF001
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
        )


def _fmt_float(value: Any) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return ""


def _fmt_probability(value: Any) -> str:
    try:
        return f"{float(value):.1%}"
    except (TypeError, ValueError):
        return ""


def _safe_timestamp(value: str) -> str:
    keep = [char for char in value if char.isdigit() or char in {"T", "Z"}]
    text = "".join(keep)
    return text or _utc_now().replace("-", "").replace(":", "")


def _forecast_datetime(forecast: dict[str, Any]) -> datetime:
    raw = str(forecast.get("generated_at") or "")
    if raw.endswith("Z"):
        raw = raw[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(raw)
    except ValueError:
        parsed = datetime.now(timezone.utc).replace(microsecond=0)
    if parsed.tzinfo is not None:
        parsed = parsed.astimezone(timezone.utc).replace(tzinfo=None)
    return parsed


def _utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
