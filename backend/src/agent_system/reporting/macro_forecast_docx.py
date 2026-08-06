"""DOCX renderer for the deterministic Helix macro forecast report."""
from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell, _Row, Table

from src.agent_system.diagnostics.input_ingestion_audit import (
    format_auditable_layer_key_signal,
    input_audit_warnings_from_input_set,
    provenance_summary_rows_from_input_set,
)
from src.agent_system.schemas.macro_forecast import (
    MacroForecastResult,
    MacroInputSignal,
    ProbabilityContribution,
    RankingContribution,
    ScenarioContribution,
    ScenarioProbabilityUpdate,
)


BODY_FONT = "Arial"
BODY_SIZE = Pt(8.5)
SMALL_SIZE = Pt(7.5)
TITLE_SIZE = Pt(26)
SUBTITLE_SIZE = Pt(18)
SECTION_SIZE = Pt(14)
SUBSECTION_SIZE = Pt(10.5)
PRIMARY_BLUE = RGBColor(0x1F, 0x4E, 0x79)
SECONDARY_BLUE = RGBColor(0x2E, 0x75, 0xB6)
BODY_GRAY = RGBColor(0x59, 0x59, 0x59)
MUTED_GRAY = RGBColor(0x80, 0x80, 0x80)
POSITIVE_GREEN = RGBColor(0x37, 0x56, 0x23)
NEGATIVE_RED = RGBColor(0xC0, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F4E79"
KEY_FILL = "EAF2F8"
ALT_ROW_FILL = "F7F9FC"
ACTIVE_PRIOR_FILL = "70AD47"
OVERRIDE_CALLOUT_FILL = "FFF2CC"
ISO_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
ISO_DATETIME_RE = re.compile(r"^\d{4}-\d{2}-\d{2}[T ].*")


def generate_macro_forecast_docx(
    result: MacroForecastResult,
    output_path: str | Path,
    *,
    debug: bool = False,
) -> Path:
    """Render a MacroForecastResult to a compact audit-friendly Word report."""

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    chart_dir = path.parent / "_macro_forecast_charts"

    document = Document()
    _style_document(document)

    _add_title(document)
    _add_metadata(document, result)
    _add_forecast_interpretation(document, result)
    _add_scenario_probabilities(document, result)
    _add_visual_summary(document, result, chart_dir)
    _add_scenario_probability_math(document, result, debug=debug)
    _add_historical_calibration(document, result, chart_dir, debug=debug)
    _add_forecast_input_set(document, result)
    _add_monetary_composite_detail(document, result)
    _add_theme_rankings(document, result)
    _add_sector_rankings(document, result)
    _add_factor_rankings(document, result)
    _add_probability_shifters(document, result)
    _add_research_priorities(document, result)
    _add_input_signal_detail(document, result)
    _add_methodology_notes(document)

    document.save(path)
    return path


def _style_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BODY_GRAY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.bold = True
        style.font.size = SECTION_SIZE if style_name == "Heading 1" else Pt(10)
        style.font.color.rgb = PRIMARY_BLUE if style_name == "Heading 1" else SECONDARY_BLUE


def _add_title(document: DocumentObject) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("HELIX INTEL")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = TITLE_SIZE
    run.font.color.rgb = PRIMARY_BLUE

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle_run = subtitle.add_run("Macro Forecast Report")
    subtitle_run.font.name = BODY_FONT
    subtitle_run.font.size = SUBTITLE_SIZE
    subtitle_run.font.color.rgb = SECONDARY_BLUE


def _add_metadata(document: DocumentObject, result: MacroForecastResult) -> None:
    primary = document.add_paragraph()
    primary.paragraph_format.space_after = Pt(2)
    primary_text = (
        f"As-of Date: {_format_date_for_display(result.asof_date)}     "
        f"Horizon: {result.horizon}     "
        f"Schema: {result.schema_version}"
    )
    primary_run = primary.add_run(primary_text)
    primary_run.bold = True
    primary_run.font.name = BODY_FONT
    primary_run.font.size = Pt(9.5)
    primary_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    secondary = document.add_paragraph()
    secondary.paragraph_format.space_after = Pt(10)
    secondary_text = f"Created At: {_text(result.created_at)}     Forecast ID: {result.id or 'unsaved'}"
    secondary_run = secondary.add_run(secondary_text)
    secondary_run.font.name = BODY_FONT
    secondary_run.font.size = Pt(8)
    secondary_run.font.color.rgb = MUTED_GRAY


def _add_forecast_interpretation(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Forecast Interpretation")
    interpretation = result.forecast_interpretation
    if interpretation is None:
        _paragraph(document, "No forecast interpretation is available.")
        return

    dominant_probability = result.scenario_probabilities.get(
        interpretation.dominant_scenario_id,
        interpretation.dominant_scenario_probability,
    )
    _kv_table(
        document,
        [
            ("Headline", interpretation.headline),
            ("Summary", interpretation.summary),
            (
                "Dominant Scenario",
                f"{_scenario_name(interpretation.dominant_scenario_id)} "
                f"({_pct(dominant_probability)})",
            ),
            ("Confidence Level", interpretation.confidence_level),
            ("Confidence Rationale", interpretation.confidence_rationale),
            ("Preferred Exposures", _join(interpretation.preferred_exposures)),
            ("Exposures To Avoid", _join(interpretation.exposures_to_avoid)),
            ("Key Tensions", _bullet_text(interpretation.key_tensions)),
        ],
    )


def _add_yaml_priors_override_callout(document: DocumentObject) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade_cell(cell, OVERRIDE_CALLOUT_FILL)
    _set_cell_text(
        cell,
        (
            "YAML PRIORS OVERRIDE ACTIVE - Engine math below shown for audit only. "
            "The Prior column is the actually-used probability distribution."
        ),
        bold=True,
        size=BODY_SIZE,
        color=RGBColor(0x7F, 0x60, 0x00),
    )
    _set_table_width(table, [9.5])
    _spacer(document, after=4)


def _add_scenario_probabilities(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Scenario Probabilities")
    calibration_by_scenario = _calibration_by_scenario(result)
    if result.probability_mode == "two_source_v1":
        headers = [
            "Scenario",
            "BVAR Soft",
            "Analogue Implied",
            "Mixed Pre-Floor",
            "Final",
            "Delta vs BVAR",
            "Floor Guard",
        ]
        mixture = result.mixture_report or {}
        per_scenario = mixture.get("per_scenario") if isinstance(mixture, dict) else {}
        rows = []
        for scenario_id, probability in sorted(
            result.scenario_probabilities.items(),
            key=lambda item: item[1],
            reverse=True,
        ):
            row = per_scenario.get(scenario_id, {}) if isinstance(per_scenario, dict) else {}
            rows.append(
                [
                    _scenario_name(scenario_id),
                    _pct(row.get("bvar_soft")),
                    _pct(row.get("analogue_implied")),
                    _pct(row.get("mixed_pre_floor")),
                    _pct(probability),
                    _pct(row.get("delta")),
                    _yes_no(bool(row.get("floor_applied", False))),
                ]
            )
    elif result.probability_mode == "yaml_priors_override":
        _add_yaml_priors_override_callout(document)
        headers = [
            "Scenario",
            "Prior (ACTIVE)",
            "Engine Pre-Floor",
            "Engine Final (audit)",
            "Engine Change",
            "Floor/Cap",
            "Top Driver",
        ]
        rows = [
            [
                _scenario_name(update.scenario_id),
                _pct(update.prior_probability),
                _pct(update.pre_floor_posterior_probability),
                _pct(update.posterior_probability),
                _pct(update.probability_change),
                _floor_cap_label(update),
                _top_driver(update),
            ]
            for update in result.scenario_updates
        ]
    elif result.probability_mode == "historically_calibrated" and calibration_by_scenario:
        headers = [
            "Scenario",
            "Prior",
            "Deterministic Posterior",
            "Historical Analogue Probability",
            "Blended Posterior",
            "Change vs Prior",
            "Floor/Cap",
            "Top Driver",
        ]
        rows = []
        for update in result.scenario_updates:
            calibration = calibration_by_scenario.get(update.scenario_id)
            if calibration is None:
                continue
            rows.append(
                [
                    _scenario_name(update.scenario_id),
                    _pct(update.prior_probability),
                    _pct(calibration.deterministic_probability),
                    _pct(calibration.historical_probability),
                    _pct(calibration.blended_probability),
                    _pct(calibration.blended_probability - update.prior_probability),
                    _floor_cap_label(update),
                    _top_driver(update),
                ]
            )
    else:
        headers = [
            "Scenario",
            "Prior",
            "Pre-Floor Posterior",
            "Final Posterior",
            "Change",
            "Floor/Cap",
            "Top Driver",
        ]
        rows = [
            [
                _scenario_name(update.scenario_id),
                _pct(update.prior_probability),
                _pct(update.pre_floor_posterior_probability),
                _pct(update.posterior_probability),
                _pct(update.probability_change),
                _floor_cap_label(update),
                _top_driver(update),
            ]
            for update in result.scenario_updates
        ]
    table = _add_table(
        document,
        headers,
        rows,
    )
    if result.probability_mode == "yaml_priors_override":
        _shade_cell(table.rows[0].cells[1], ACTIVE_PRIOR_FILL)
        _set_cell_text(table.rows[0].cells[1], "Prior (ACTIVE)", bold=True, size=SMALL_SIZE, color=WHITE)
        _set_table_width(table, [1.75, 0.8, 1.0, 1.05, 0.75, 0.9, 2.45])
    elif result.probability_mode == "two_source_v1":
        _set_table_width(table, [1.45, 0.8, 0.9, 0.95, 0.75, 0.8, 0.75])
    elif len(headers) == 8:
        _set_table_width(table, [1.45, 0.65, 0.9, 1.0, 0.9, 0.75, 0.8, 2.15])
    else:
        _set_table_width(table, [1.75, 0.7, 1.1, 1.0, 0.7, 0.9, 2.6])


def _add_visual_summary(document: DocumentObject, result: MacroForecastResult, chart_dir: Path) -> None:
    _section(document, "Visual Summary")
    for title, chart_path in [
        ("Scenario Probability Chart", _scenario_probability_chart(result, chart_dir)),
        ("Theme Macro Support Chart", _theme_support_chart(result, chart_dir)),
        ("Historical Macro Return Profile Chart", _macro_forward_return_chart(result, chart_dir)),
    ]:
        if chart_path is None:
            _paragraph(document, f"{title}: insufficient data for chart.", compact=True)
            continue
        _paragraph(document, title, bold=True, compact=True)
        _insert_chart(document, chart_path)


def _add_scenario_probability_math(document: DocumentObject, result: MacroForecastResult, *, debug: bool = False) -> None:
    _section(document, "Scenario Probability Math")
    if result.probability_mode == "two_source_v1":
        mixture = result.mixture_report or {}
        evidence = mixture.get("evidence") if isinstance(mixture, dict) else {}
        provenance = result.bvar_provenance or {}
        _kv_table(
            document,
            [
                ("Probability mode", "two_source_v1"),
                ("Combination", str(mixture.get("combination") or "linear_mixture")),
                ("Alpha", _decimal(mixture.get("alpha"), places=3)),
                ("Alpha effective", _decimal(mixture.get("alpha_effective"), places=3)),
                ("Analogue trailing share", _pct(mixture.get("s"))),
                ("Analogue base rate", _pct(mixture.get("b"))),
                ("Evidence state", str((evidence or {}).get("current_state") or mixture.get("abstention_state") or "n/a")),
                ("Stress advisory", _yes_no(bool(mixture.get("stress_advisory", False)))),
                ("BVAR artifact", str(provenance.get("path") or "n/a")),
                ("BVAR generated_at", str(provenance.get("generated_at") or "n/a")),
                ("BVAR handoff fingerprint", str(provenance.get("handoff_fingerprint") or "n/a")),
            ],
        )
        limitations = provenance.get("model_limitations") or {}
        if limitations:
            _subsection(document, "BVAR Model Limitations")
            _kv_table(document, [(str(key), str(value)) for key, value in limitations.items()])
        return
    for update in result.scenario_updates:
        audit = update.math_audit
        _subsection(document, _scenario_name(update.scenario_id))
        if audit is None:
            _paragraph(document, "No math audit is available for this scenario.")
            continue

        _kv_table(
            document,
            [
                ("Prior", _pct(audit.prior_probability)),
                ("Prior score/log score", _signed(audit.prior_logit_or_log_score)),
                ("Raw score before softmax", _signed(audit.raw_score_before_softmax)),
                ("Pre-floor posterior", _pct(audit.pre_floor_posterior_probability)),
                ("Final posterior", _pct(audit.final_posterior_probability)),
                ("Floor applied", f"{audit.floor_applied} ({_pct(audit.floor_value)})"),
                ("Cap applied", f"{audit.cap_applied} ({_pct(audit.cap_value)})"),
            ],
        )

        _paragraph(document, "Contributor Breakdown", bold=True, compact=True)
        _contributor_breakdown_table(document, update)
        if not _contributions_for_roles(update, {"raw_component"}):
            _paragraph(document, "No raw component contributors were used for this scenario.", compact=True)
        if debug:
            for title, roles in [
                ("Layer Summary Contributors", {"layer_summary"}),
                ("Raw Component Contributors", {"raw_component"}),
                ("Composite Contributors", {"composite"}),
                ("Regime Driver Contributors", {"regime_driver"}),
                ("Scenario Falsifier Contributors", {"scenario_falsifier"}),
            ]:
                items = _contributions_for_roles(update, roles)
                if items:
                    _paragraph(document, title, bold=True, compact=True)
                    _contributor_table(document, items)


def _add_historical_calibration(
    document: DocumentObject,
    result: MacroForecastResult,
    chart_dir: Path,
    *,
    debug: bool = False,
) -> None:
    _section(document, "Historical Analogue Calibration")
    calibration = result.historical_calibration
    if result.probability_mode == "two_source_v1":
        _paragraph(
            document,
            "Legacy rolling historical calibration is retired in the runner. Directional analogue evidence enters through the two-source mixture.",
            compact=True,
        )
        return
    if calibration is None:
        _paragraph(document, "Historical calibration is disabled.", compact=True)
        return

    _kv_table(
        document,
        [
            ("Enabled", _yes_no(calibration.enabled)),
            ("Method", calibration.method),
            ("Analogue version", calibration.analogue_version),
            ("Conditions summary", calibration.conditions_summary or "n/a"),
            ("Analogues", f"{calibration.n_analogues}"),
            ("Unique / pooled", f"{calibration.n_unique_analogues or 'n/a'} / {calibration.n_pooled or 'n/a'}"),
            ("Confidence", _decimal(calibration.confidence, places=2)),
            ("Warnings", _bullet_text(calibration.warnings)),
        ],
    )

    if calibration.shock_window_diagnostics:
        diagnostics = calibration.shock_window_diagnostics
        _subsection(document, "Shock Window Diagnostics")
        windows = diagnostics.get("windows") or []
        window_text = "; ".join(
            f"{window.get('name', 'shock')} "
            f"{_format_date_for_display(window.get('start_date'))} to "
            f"{_format_date_for_display(window.get('end_date'))}"
            for window in windows
            if isinstance(window, dict)
        ) or "none"
        _kv_table(
            document,
            [
                ("Mode", str(diagnostics.get("mode") or "n/a")),
                ("Shock windows applied", window_text),
                ("Rows excluded by horizon", _horizon_count_text(diagnostics.get("rows_excluded_by_horizon") or {})),
                ("Analogue dates excluded from scenario mapping", _join(diagnostics.get("scenario_mapping_excluded_dates") or [])),
                ("Scenario mapping horizon", str(diagnostics.get("scenario_mapping_horizon") or "n/a")),
                ("Historical probabilities changed", _yes_no(bool(diagnostics.get("historical_probabilities_changed")))),
            ],
        )

    if calibration.analogue_version == "v2_detailed":
        diagnostics = calibration.detailed_analogue_diagnostics or {}
        _subsection(document, "Detailed Analogue Match Quality")
        _kv_table(
            document,
            [
                ("V1 broad-state weight", _decimal(diagnostics.get("v1_weight"), places=2)),
                ("V2 detailed-input weight", _decimal(diagnostics.get("v2_weight"), places=2)),
                ("Candidate pool size", diagnostics.get("candidate_pool_n") or "n/a"),
                ("Effective sample size", _decimal(diagnostics.get("effective_sample_size"), places=2)),
                ("Adjusted deterministic weight", _decimal(diagnostics.get("adjusted_deterministic_weight"), places=2)),
                ("Adjusted historical weight", _decimal(diagnostics.get("adjusted_historical_weight"), places=2)),
                ("Average detailed similarity", _decimal(diagnostics.get("average_detailed_similarity"), places=2)),
                ("Average blended similarity", _decimal(diagnostics.get("average_blended_similarity"), places=2)),
                ("Strongest groups", _join(diagnostics.get("strongest_match_groups") or [])),
                ("Weakest groups", _join(diagnostics.get("weakest_match_groups") or [])),
                ("Missing important features", _join(diagnostics.get("missing_important_features") or [])),
            ],
        )
        _add_table(
            document,
            ["Group", "Avg Similarity", "Features Used", "Features Missing", "Coverage", "Used Feature IDs", "Missing Feature IDs", "Interpretation"],
            _detailed_match_quality_rows(diagnostics.get("group_similarity_summary") or {}),
        )

    _subsection(document, "Historical Macro Forward Return Stats")
    macro_stats = calibration.macro_forward_return_stats or {
        horizon: calibration.forward_return_stats[horizon]
        for horizon in ["21d", "63d", "126d", "252d"]
        if horizon in calibration.forward_return_stats
    }
    _add_table(
        document,
        ["Horizon", "N", "Weight Sum", "Median", "Mean", "% Positive", "P10", "P25", "P75", "P90", "Worst", "Best"],
        _forward_stat_rows(macro_stats, ["21d", "63d", "126d", "252d"]),
    )
    chart_path = _macro_forward_return_chart(result, chart_dir)
    if chart_path is not None:
        _insert_chart(document, chart_path)
    _paragraph(
        document,
        "Longer horizons naturally have lower sample sizes because recent historical rows do not yet have full forward-return windows, especially 1Y / 252D.",
        compact=True,
    )

    tactical_stats = calibration.tactical_forward_return_stats or {
        horizon: calibration.forward_return_stats[horizon]
        for horizon in ["1d", "5d", "10d"]
        if horizon in calibration.forward_return_stats
    }
    if tactical_stats:
        _subsection(document, "Historical Tactical Forward Return Stats")
        _add_table(
            document,
            ["Horizon", "N", "Weight Sum", "Median", "Mean", "% Positive", "P10", "P25", "P75", "P90", "Worst", "Best"],
            _forward_stat_rows(tactical_stats, ["1d", "5d", "10d"]),
        )

    if calibration.risk_profile:
        _subsection(document, "Historical Risk Profile by Horizon")
        rows = _risk_profile_by_horizon_rows(calibration.risk_profile)
        if rows:
            _add_table(
                document,
                ["Horizon", "Win Rate", "Median Up", "Median Down", "Expected Value", "P10 Return", "P90 Return", "Worst Return"],
                rows,
            )
            risk_chart = _risk_profile_chart(result, chart_dir)
            if risk_chart is not None:
                _insert_chart(document, risk_chart)
        else:
            _paragraph(document, "Macro-horizon drawdown/upside values are unavailable; return-distribution risk is shown in the forward-return stats.", compact=True)
        tactical_rows = _tactical_risk_rows(calibration.risk_profile)
        if tactical_rows:
            _subsection(document, "Historical Tactical Risk Snapshot")
            _add_table(document, ["Metric", "Value"], tactical_rows)

    _subsection(document, "Top Analogue Dates")
    rows = [
        [
            analogue.date,
            _signed(analogue.composite_weight),
            _compact_env(analogue.environment),
            _signed(analogue.score_total),
            _signed(analogue.vix_level),
            str(analogue.sectors_green) if analogue.sectors_green is not None else "-",
            _signed(analogue.forward_returns.get("63d")),
            _signed(analogue.forward_returns.get("21d")),
            analogue.mapped_scenario_id or "-",
            analogue.mapping_rationale_short or analogue.mapping_tag or "-",
            _signed(analogue.v1_similarity),
            _signed(analogue.detailed_similarity),
            _signed(analogue.blended_similarity),
            _join(analogue.strongest_matching_groups),
            _join(analogue.weakest_matching_groups),
        ]
        for analogue in calibration.top_analogues[:10]
    ]
    _add_table(
        document,
        ["Date", "Weight", "Env", "Score", "VIX", "Breadth", "63D Fwd", "21D Fwd", "Scenario", "Map Tag", "V1", "V2", "Blend", "Strong Groups", "Weak Groups"],
        rows,
    )
    if debug:
        _subsection(document, "Analogue Mapping Detail")
        _add_table(
            document,
            ["Date", "Full Mapping Rationale"],
            [
                [analogue.date, analogue.mapping_rationale_full or analogue.mapping_rationale or "-"]
                for analogue in calibration.top_analogues[:10]
            ],
        )

    _subsection(document, "Scenario Calibration")
    rows = [
        [
            _scenario_name(item.scenario_id),
            _pct(item.deterministic_probability),
            _pct(item.historical_probability),
            _pct(item.blended_probability),
            _pct(item.analog_effect),
            str(item.n_supporting_analogues),
            _decimal(item.confidence, places=2),
            item.rationale,
        ]
        for item in calibration.scenario_calibrations
    ]
    _add_table(
        document,
        [
            "Scenario",
            "Deterministic Probability",
            "Historical Probability",
            "Blended Probability",
            "Analog Effect",
            "Supporting Analogues",
            "Confidence",
            "Rationale",
        ],
        rows,
    )
    scenario_chart = _scenario_calibration_chart(result, chart_dir)
    if scenario_chart is not None:
        _insert_chart(document, scenario_chart)

    _subsection(document, "Historical Calibration Methodology")
    for note in calibration.methodology_notes:
        _paragraph(document, note, compact=True)


def _add_forecast_input_set(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Forecast Input Set")
    _paragraph(document, "Monitoring — no probability impact.", compact=True)
    current_regime_yaml_path = (result.outputs or {}).get("current_regime_yaml_path")
    if current_regime_yaml_path:
        _paragraph(document, f"Thematic agent current-regime YAML saved to: {current_regime_yaml_path}", compact=True)
    input_set = result.forecast_input_set
    if input_set is None:
        _paragraph(document, "ForecastInputSet is unavailable; showing the legacy flat input signal list.", compact=True)
        _add_legacy_input_signal_table(document, result.input_signals)
        return

    diagnostic_path = Path("data/agent_system/diagnostics") / f"input_matrix_{result.asof_date}.csv"
    ingestion_audit_path = Path("data/agent_system/diagnostics") / f"input_ingestion_audit_{result.asof_date}.csv"
    _paragraph(document, f"Full input diagnostic matrix saved to: {diagnostic_path}", compact=True)
    _paragraph(document, f"Full input ingestion audit saved to: {ingestion_audit_path}", compact=True)

    input_audit_warnings = input_audit_warnings_from_input_set(input_set, horizon=result.horizon)
    if input_audit_warnings:
        _subsection(document, "Input Audit Warnings")
        for warning in input_audit_warnings:
            _paragraph(document, warning, compact=True)

    input_warnings = _forecast_input_warnings(result)
    if input_warnings:
        _subsection(document, "Forecast Input Warnings")
        for warning in input_warnings:
            _paragraph(document, warning, compact=True)

    _add_input_provenance_summary(document, result)

    _subsection(document, "Layer Summary Signals")
    rows = [
        [
            signal.parent_layer or signal.category,
            _signal_reading(signal.current_value),
            f"{signal.signal}/{signal.trend}",
            f"{_decimal(signal.confidence, places=2)}/{signal.data_quality}",
            _yes_no(signal.used_in_probability_update),
            format_auditable_layer_key_signal(input_set, signal, horizon=result.horizon),
            _raw_components_for_layer(result, signal.parent_layer),
        ]
        for signal in input_set.layer_summary_signals
    ]
    _add_table(
        document,
        ["Layer", "Score", "Status", "Confidence/Data Quality", "Probability Impact?", "Key Signals", "Raw Components Attached"],
        rows,
    )

    _subsection(document, "Raw Input Coverage Summary")
    coverage_rows, coverage_totals = _raw_input_coverage_rows(input_set)
    _add_table(
        document,
        ["Parent Layer", "Expected Raw Inputs", "Available Raw Inputs", "Used in Retired Deterministic Math", "Used in Historical Similarity", "Missing Inputs"],
        coverage_rows,
    )
    _subsection(document, "Raw Input Coverage Totals")
    _add_table(
        document,
        ["Metric", "Count"],
        [[key, value] for key, value in coverage_totals.items()],
    )

    _subsection(document, "Raw Component Signals")
    raw_signals = sorted(
        input_set.raw_component_signals,
        key=lambda signal: float(signal.signal_strength or 0.0),
        reverse=True,
    )
    raw_note_needed = len(raw_signals) > 25
    rows = [
        [
            signal.name,
            signal.source_object or "-",
            signal.parent_layer or "-",
            _signal_reading(signal.raw_value),
            signal.level_status or signal.signal,
            signal.trend_status or signal.trend,
            _yes_no(signal.used_in_probability_update),
            _yes_no(signal.used_in_historical_similarity),
            _top_scenario_effects(signal),
        ]
        for signal in raw_signals[:25]
    ]
    _add_table(
        document,
        ["Input", "Source Object", "Parent Layer", "Raw Value", "Status", "Trend", "Probability Impact?", "Used in Analogues?", "Top Scenario Effects"],
        rows,
    )
    if raw_note_needed:
        _paragraph(document, "Full raw signal list available in JSON.", compact=True)

    _subsection(document, "Composite Signals")
    rows = [
        [
            signal.name,
            signal.parent_layer or "-",
            signal.signal,
            _decimal(signal.confidence, places=2),
            _yes_no(signal.used_in_probability_update),
            _join(signal.child_signal_ids),
            _top_scenario_effects(signal),
        ]
        for signal in input_set.composite_signals
    ]
    _add_table(
        document,
        ["Composite", "Parent Layer", "Signal", "Confidence", "Probability Impact?", "Child Signals", "Scenario Effects"],
        rows,
    )

    _subsection(document, "Market/Tape Signals")
    rows = [
        [
            signal.name,
            _signal_reading(signal.raw_value),
            signal.level_status or signal.signal,
            _decimal(signal.confidence, places=2),
            _yes_no(signal.used_in_probability_update),
            signal.transformation_method or "-",
            _top_scenario_effects(signal),
        ]
        for signal in input_set.market_tape_signals
    ]
    _add_table(
        document,
        ["Input", "Raw Value", "Status", "Confidence", "Probability Impact?", "Horizon/Dedupe Effect", "Top Scenario Effects"],
        rows,
    )

    _subsection(document, "Regime-Specific Drivers")
    rows = [
        [
            signal.name,
            signal.parent_layer or "-",
            signal.signal,
            _decimal(signal.confidence, places=2),
            _join(signal.active_only_in_regime_ids),
            _top_scenario_effects(signal),
        ]
        for signal in input_set.regime_driver_signals
    ]
    _add_table(
        document,
        ["Driver", "Parent Layer", "Signal", "Confidence", "Active Regime", "Scenario Effects"],
        rows,
    )

    _subsection(document, "Scenario Falsifiers")
    rows = [
        [
            signal.name,
            _join(signal.related_scenario_ids),
            signal.signal,
            _decimal(signal.confidence, places=2),
            _signal_reading(signal.current_value),
            _top_scenario_effects(signal),
        ]
        for signal in input_set.scenario_falsifier_signals
    ]
    _add_table(
        document,
        ["Falsifier", "Related Scenarios", "Signal", "Confidence", "Current Value", "Scenario Effects"],
        rows,
    )

    _subsection(document, "Dedupe / Weighting Notes")
    for note in input_set.methodology_notes:
        _paragraph(document, note, compact=True)


def _add_input_provenance_summary(document: DocumentObject, result: MacroForecastResult) -> None:
    input_set = result.forecast_input_set
    rows = provenance_summary_rows_from_input_set(input_set, asof_date=result.asof_date, horizon=result.horizon)
    if not rows:
        return
    _subsection(document, "Input Provenance Summary")
    table_rows = []
    for row in rows:
        source_parts = [
            row.get("source_name") or row.get("provider"),
            row.get("source_object"),
            row.get("source_field"),
        ]
        source = ".".join(str(part) for part in source_parts if part)
        table_rows.append(
            [
                row.get("display_label") or row.get("input_id"),
                _signal_reading(row.get("value")),
                source or "-",
                row.get("observed_date") or "?",
                row.get("lookback_window") or "?",
                row.get("freshness_status") or "unknown",
                row.get("audit_status") or "unknown",
            ]
        )
    _add_table(
        document,
        ["Input", "Value", "Source", "Observed Date", "Lookback", "Freshness", "Audit Status"],
        table_rows,
    )


def _add_legacy_input_signal_table(document: DocumentObject, signals: Sequence[MacroInputSignal]) -> None:
    rows = [
        [
            signal.name,
            signal.category,
            _signal_reading(signal.current_value),
            signal.signal,
            signal.trend,
            _decimal(signal.confidence, places=2),
            signal.data_quality,
            _yes_no(signal.used_in_probability_update),
            _yes_no(signal.display_only),
            signal.parent_signal_id or ("composite" if signal.child_signal_ids else "-"),
            signal.exclusion_reason or "-",
        ]
        for signal in signals
    ]
    _add_table(
        document,
        [
            "Signal",
            "Category",
            "Reading",
            "Signal",
            "Trend",
            "Confidence",
            "Quality",
            "Probability Impact?",
            "Display Only?",
            "Parent / Composite",
            "Exclusion Reason",
        ],
        rows,
    )


def _forecast_input_warnings(result: MacroForecastResult) -> list[str]:
    input_set = result.forecast_input_set
    if input_set is None:
        return ["ForecastInputSet unavailable; raw input coverage could not be audited."]
    coverage = getattr(input_set, "raw_input_coverage", {}) or {}
    groups = coverage.get("groups") if isinstance(coverage, dict) else None
    volatility = groups.get("volatility") if isinstance(groups, dict) else None
    warnings: list[str] = []
    if isinstance(coverage, dict):
        warnings.extend(str(item) for item in (coverage.get("warnings") or []))
    if not isinstance(volatility, dict):
        warnings.append("Volatility raw input coverage unavailable.")
    else:
        missing = [str(item) for item in volatility.get("missing") or []]
        if missing:
            warnings.append(f"Volatility raw input coverage incomplete; missing: {', '.join(missing)}.")
    totals = coverage.get("totals") if isinstance(coverage, dict) else None
    if isinstance(totals, dict):
        expected = float(totals.get("total_raw_signals_expected") or 0.0)
        available = float(totals.get("total_raw_signals_available") or 0.0)
        if expected > 0 and available / expected < 0.40:
            warnings.append("Raw input coverage is low; forecast is relying primarily on layer summaries and regime drivers.")

    available_raw = [
        signal
        for signal in input_set.raw_component_signals
        if signal.raw_value is not None and not signal.display_only
    ]
    raw_contributions_used = [
        contribution
        for update in result.scenario_updates
        for contribution in update.contributions
        if contribution.source_role == "raw_component" and abs(contribution.adjusted_contribution) > 0
    ]
    if result.probability_mode != "two_source_v1" and available_raw and not raw_contributions_used:
        warnings.append("Raw component signals were available but did not affect deterministic scenario math. Check dedupe/config logic.")
    return list(dict.fromkeys(warnings))


def _add_monetary_composite_detail(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Monetary Composite Detail")
    composite = next(
        (signal for signal in result.input_signals if signal.input_id == "monetary_policy_composite"),
        None,
    )
    if composite is None:
        _paragraph(document, "No monetary composite signal is present.")
        return

    _kv_table(
        document,
        [
            ("Composite signal", composite.signal),
            ("Composite confidence", _decimal(composite.confidence, places=2)),
            ("Probability impact", _yes_no(composite.used_in_probability_update)),
            ("Display only", _yes_no(composite.display_only)),
            ("Current value", _signal_reading(composite.current_value)),
            ("Composite method", composite.composite_method or "n/a"),
            ("Notes", composite.notes),
        ],
    )

    child_ids = set(composite.child_signal_ids)
    children = [signal for signal in result.input_signals if signal.input_id in child_ids]
    if children:
        _paragraph(
            document,
            "Component signals are included for monitoring and excluded from two_source_v1 probability math.",
            compact=True,
        )
        rows = [
            [
                child.name,
                child.signal,
                child.trend,
                _decimal(child.confidence, places=2),
                _yes_no(child.display_only),
                _yes_no(child.used_in_probability_update),
                child.exclusion_reason or "-",
            ]
            for child in children
        ]
        _add_table(
            document,
            ["Component", "Signal", "Trend", "Confidence", "Display Only?", "Probability Impact?", "Exclusion Reason"],
            rows,
        )

    impacts = [
        f"{impact.scenario_id}: {impact.direction} {impact.strength:.2f}"
        for impact in composite.affected_scenarios
    ]
    _paragraph(document, f"Assigned scenario impacts: {_join(impacts)}", compact=True)


def _add_theme_rankings(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Theme Rankings - Macro Support Score")
    _paragraph(
        document,
        "Theme rankings in the macro forecast are based only on macro/scenario support. The score is the "
        "probability-weighted scenario exposure of each theme. Crowding, valuation, narrative maturity, "
        "consensus gap, and ticker-level dispersion are intentionally excluded here and evaluated by downstream agents.",
        compact=True,
    )
    _paragraph(
        document,
        "Deprecated overlay fields may exist in historical records but are not used in current macro ranking.",
        compact=True,
    )
    rows = [
        [
            str(index),
            theme.label,
            _signed(theme.macro_support_score),
            _join(theme.best_scenarios),
            _join(theme.worst_scenarios),
            _theme_contribution_summary(theme),
            "Macro-supported research direction; downstream agents evaluate crowding, valuation, narrative maturity, consensus gap, and ticker quality.",
        ]
        for index, theme in enumerate(result.theme_rankings, 1)
    ]
    table = _add_table(
        document,
        [
            "Rank",
            "Theme",
            "Macro Support Score",
            "Best Scenarios",
            "Worst Scenarios",
            "Scenario Contribution Breakdown",
            "Interpretation",
        ],
        rows,
    )
    _set_table_width(table, [0.35, 1.55, 0.85, 1.35, 1.35, 2.65, 2.25])

    _subsection(document, "Theme Macro Support Math")
    for theme in result.theme_rankings[:5]:
        _paragraph(document, f"{theme.label}: {_signed(theme.macro_support_score)}", bold=True, compact=True)
        rows = [
            [
                contribution.scenario_label,
                _pct(contribution.scenario_probability),
                _signed(contribution.theme_exposure_score),
                _signed(contribution.contribution),
                contribution.rationale or "-",
            ]
            for contribution in sorted(
                theme.scenario_contributions,
                key=lambda item: abs(item.contribution),
                reverse=True,
            )
        ]
        rows.append(
            [
                "Net",
                "-",
                "-",
                _signed(theme.net_macro_support_score),
                (
                    f"Positive total {_signed(theme.positive_contribution_total)}; "
                    f"negative total {_signed(theme.negative_contribution_total)}."
                ),
            ]
        )
        _add_table(document, ["Scenario", "Probability", "Exposure", "Contribution", "Rationale"], rows)


def _add_sector_rankings(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Sector & Instrument Rankings")
    rows = [
        [
            str(index),
            ranking.ticker or ranking.label or ranking.item_id or "-",
            _signed(ranking.score),
            _ranking_driver_summary(ranking.contributions),
            _join(ranking.formula_notes),
        ]
        for index, ranking in enumerate(result.sector_rankings, 1)
    ]
    _add_table(document, ["Rank", "Ticker", "Score", "Top Drivers", "Formula Notes"], rows)


def _add_factor_rankings(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Factor Rankings")
    rows = [
        [
            str(index),
            ranking.label or ranking.factor_id or ranking.item_id or "-",
            _signed(ranking.score),
            _ranking_driver_summary(ranking.contributions),
            _join(ranking.formula_notes),
        ]
        for index, ranking in enumerate(result.factor_rankings, 1)
    ]
    _add_table(document, ["Rank", "Factor", "Score", "Top Drivers", "Formula Notes"], rows)


def _add_probability_shifters(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Probability Shifters / Watchlist")
    if result.probability_mode == "two_source_v1":
        _paragraph(
            document,
            "Deterministic probability shifters are retired; monitoring signals and falsifiers remain display-only.",
            compact=True,
        )
        return
    for shifter in result.probability_shifters:
        _subsection(document, _scenario_name(shifter.scenario_id))
        _kv_table(
            document,
            [
                ("Current probability", _pct(shifter.current_probability)),
                ("Would increase if", _bullet_text(shifter.would_increase_probability_if)),
                ("Would decrease if", _bullet_text(shifter.would_decrease_probability_if)),
                ("Key inputs to watch", _join(shifter.key_inputs_to_watch)),
                ("Floor/cap note", shifter.floor_or_cap_note or "-"),
            ],
        )


def _add_research_priorities(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Recommended Research Priorities")
    for priority in result.recommended_research_priorities:
        _subsection(document, f"{priority.priority_rank}. {priority.theme}")
        _kv_table(
            document,
            [
                ("Theme", priority.theme),
                ("Edge hypothesis", priority.edge_hypothesis),
                ("Expected edge decay", str(priority.expected_edge_decay)),
                ("Rationale", priority.rationale),
                ("Research sub-questions", _bullet_text(priority.sub_questions)),
                ("Supporting evidence", _evidence_text(priority.supporting_evidence)),
            ],
        )


def _add_input_signal_detail(document: DocumentObject, result: MacroForecastResult) -> None:
    _section(document, "Input Signal Detail")
    for signal in result.input_signals:
        _subsection(document, signal.name)
        scenario_effects = [
            f"{impact.scenario_id}: {impact.direction} {impact.strength:.2f} - {impact.rationale}"
            for impact in signal.affected_scenarios
        ]
        theme_effects = [
            f"{impact.theme_id}: {impact.direction} {impact.strength:.2f} - {impact.rationale}"
            for impact in signal.affected_themes
        ]
        _kv_table(
            document,
            [
                ("Signal", f"{signal.signal} / {signal.trend} / confidence {signal.confidence:.2f}"),
                ("Notes", signal.notes or "-"),
                ("Scenario effects", _bullet_text(scenario_effects)),
                ("Theme effects", _bullet_text(theme_effects)),
            ],
        )


def _add_methodology_notes(document: DocumentObject) -> None:
    _section(document, "Methodology Notes")
    formulas = [
        "Layer summaries are generated from Helix regime layer scores.",
        "Raw components are generated from the underlying RegimeInputs fields.",
        "Market/tape signals are generated from MarketState.",
        "Regime drivers and scenario falsifiers are generated from active regime context and scenario definitions.",
        "In hybrid mode, layer summaries provide base scenario impacts and raw components act as modifiers within the same parent layer.",
        "Dedupe caps prevent one layer from dominating through many correlated inputs.",
        "Market/tape signals are downweighted for longer horizons.",
        "raw_component_contribution = direction_sign × base_strength × confidence × signal_multiplier × horizon_weight × dedupe_weight",
        "layer_summary_contribution = direction_sign × base_strength × confidence × layer_summary_base_weight",
        "final_layer_group_contribution = layer_summary_contribution + capped_sum(raw_component_modifiers)",
        "raw_score_s = prior_score_s + Σ input_contribution_i,s",
        "input_contribution_i,s = direction_sign × base_strength_i,s × confidence_i × signal_multiplier_i",
        "pre_floor_probability_s = softmax(raw_score_s)",
        "final_probability_s = apply_floors_and_caps(pre_floor_probability_s)",
        "macro_support_score_t = Σ scenario_probability_s × theme_exposure_score_t,s",
        "theme_contribution_t,s = scenario_probability_s × theme_exposure_score_t,s",
        "ranking_score_t = macro_support_score_t",
        "sector_score = Σ theme_macro_support_score_t × sector_theme_weight_t",
        "factor_score = Σ theme_macro_support_score_t × factor_theme_weight_t",
    ]
    for formula in formulas:
        _paragraph(document, formula, compact=True)
    _paragraph(
        document,
        "Macro forecast theme rankings intentionally exclude valuation, crowding, narrative maturity, "
        "consensus gap, and ticker-level quality. Those are evaluated by downstream research agents.",
        compact=True,
    )
    _paragraph(
        document,
        "Monetary component signals are displayed as inputs but excluded from probability math when "
        "monetary_policy_composite is enabled.",
        compact=True,
    )


def _horizon_label(horizon: str) -> str:
    return {
        "1d": "1D",
        "5d": "5D",
        "10d": "10D",
        "21d": "1M / 21D",
        "63d": "3M / 63D",
        "126d": "6M / 126D",
        "252d": "1Y / 252D",
    }.get(horizon, horizon.upper())


def _forward_stat_rows(stats_by_horizon, horizons: Sequence[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for horizon in horizons:
        stats = stats_by_horizon.get(horizon)
        if stats is None:
            continue
        rows.append(
            [
                _horizon_label(horizon),
                str(stats.n),
                _decimal(stats.weight_sum),
                _signed(stats.median),
                _signed(stats.mean),
                _decimal(stats.pct_positive),
                _signed(stats.p10),
                _signed(stats.p25),
                _signed(stats.p75),
                _signed(stats.p90),
                _signed(stats.worst),
                _signed(stats.best),
            ]
        )
    return rows


def _risk_profile_by_horizon_rows(risk_profile: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for horizon in ["21d", "63d", "126d", "252d"]:
        row_values = [
            risk_profile.get(f"win_rate_{horizon}"),
            risk_profile.get(f"median_up_{horizon}"),
            risk_profile.get(f"median_down_{horizon}"),
            risk_profile.get(f"expected_value_{horizon}"),
            risk_profile.get(f"p10_forward_return_{horizon}"),
            risk_profile.get(f"p90_forward_return_{horizon}"),
            risk_profile.get(f"worst_forward_return_{horizon}"),
        ]
        if all(value is None for value in row_values):
            continue
        rows.append(
            [
                _horizon_label(horizon),
                _decimal(row_values[0]),
                _signed(row_values[1]),
                _signed(row_values[2]),
                _signed(row_values[3]),
                _signed(row_values[4]),
                _signed(row_values[5]),
                _signed(row_values[6]),
            ]
        )
    return rows


def _tactical_risk_rows(risk_profile: dict) -> list[list[str]]:
    mapping = [
        ("Median 5D drawdown", risk_profile.get("median_max_drawdown_5d")),
        ("Median 5D upside", risk_profile.get("median_max_upside_5d")),
        ("Reward/risk ratio", risk_profile.get("reward_risk_ratio")),
    ]
    return [[label, _signed(value)] for label, value in mapping if value is not None]


def _detailed_match_quality_rows(summary: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for group, values in sorted(summary.items()):
        coverage = values.get("coverage") if isinstance(values, dict) else None
        avg_similarity = values.get("avg_similarity") if isinstance(values, dict) else None
        features_used = values.get("features_used") if isinstance(values, dict) else None
        features_missing = values.get("features_missing") if isinstance(values, dict) else None
        top_used = values.get("top_features_used") if isinstance(values, dict) else None
        top_missing = values.get("top_features_missing") if isinstance(values, dict) else None
        if coverage is None:
            interpretation = "Coverage unavailable"
        elif coverage >= 0.75:
            interpretation = "Good coverage"
        elif coverage >= 0.40:
            interpretation = "Partial coverage"
        else:
            interpretation = "Sparse coverage"
        rows.append(
            [
                str(group),
                _decimal(avg_similarity, places=2),
                str(features_used if features_used is not None else "n/a"),
                str(features_missing if features_missing is not None else "n/a"),
                _decimal(coverage, places=2),
                _join(top_used or []),
                _join(top_missing or []),
                interpretation,
            ]
        )
    return rows


def _raw_input_coverage_rows(input_set) -> tuple[list[list[str]], dict[str, int]]:
    expected_by_layer = {
        "Monetary": ["net_liquidity", "net_liquidity_z", "nfci", "nfci_inverted", "m2_growth_yoy", "fci_z"],
        "Credit": ["hy_spread_level", "hy_spread_z", "hy_spread_chg_4w", "ig_spread_level", "ig_spread_z", "hyg_tlt_ratio_z"],
        "Volatility": ["vix_level", "vix_z_20d", "vix_term_slope", "vvix_level", "vvix_z", "put_call_ratio", "skew_index"],
        "Breadth": ["pct_above_200d", "new_highs_minus_lows_z", "sectors_green", "rsp_vs_spy_z", "adl_slope"],
        "Positioning": ["dealer_gamma_z", "put_call_5d_ma", "aaii_bull_minus_bear", "cot_net_large_spec_z", "equity_etf_flow_z"],
        "Market/Tape": ["spy_clv", "spy_range_pct", "spy_vol_z_20d", "volume_confirmation", "spy_above_vwap", "spy_above_prev_close", "sector_dispersion", "sectors_green", "leadership_top3", "rsp_minus_spy", "iwm_minus_spy", "hyg_minus_tlt", "qqq_minus_spy"],
        "Rates/FX": [],
        "Commodities/Oil": [],
        "Earnings/Theme": [],
    }
    parent_to_label = {
        "monetary": "Monetary",
        "credit": "Credit",
        "volatility": "Volatility",
        "breadth": "Breadth",
        "positioning": "Positioning",
        "market_state": "Market/Tape",
        "rates_fx": "Rates/FX",
        "commodities": "Commodities/Oil",
        "earnings": "Earnings/Theme",
    }
    signals_by_label: dict[str, list[MacroInputSignal]] = {label: [] for label in expected_by_layer}
    for signal in list(input_set.raw_component_signals) + list(input_set.market_tape_signals) + list(input_set.theme_specific_signals):
        label = parent_to_label.get(str(signal.parent_layer or ""), None)
        if label is None and signal.input_scope == "market_tape":
            label = "Market/Tape"
        if label in signals_by_label:
            signals_by_label[label].append(signal)

    rows: list[list[str]] = []
    total_expected = total_available = total_prob = total_hist = total_display_only = total_missing = 0
    for label, expected in expected_by_layer.items():
        signals = signals_by_label.get(label, [])
        available_ids = [signal.input_id for signal in signals if signal.raw_value is not None]
        expected_set = set(expected)
        missing = [field for field in expected if field not in set(available_ids)]
        if not expected and signals:
            expected = available_ids
            expected_set = set(expected)
            missing = []
        used_prob = [
            signal.input_id
            for signal in signals
            if signal.used_in_probability_update and (not expected_set or signal.input_id in expected_set)
        ]
        used_hist = [
            signal.input_id
            for signal in signals
            if signal.used_in_historical_similarity and (not expected_set or signal.input_id in expected_set)
        ]
        display_only = [
            signal.input_id
            for signal in signals
            if signal.display_only and (not expected_set or signal.input_id in expected_set)
        ]
        total_expected += len(expected)
        total_available += len(available_ids)
        total_prob += len(used_prob)
        total_hist += len(used_hist)
        total_display_only += len(display_only)
        total_missing += len(missing)
        rows.append(
            [
                label,
                str(len(expected)),
                str(len(available_ids)),
                str(len(used_prob)),
                str(len(used_hist)),
                _join(missing[:8]),
            ]
        )
    totals = {
        "total_raw_signals_expected": total_expected,
        "total_raw_signals_available": total_available,
        "total_raw_signals_used_in_probability_update": total_prob,
        "total_raw_signals_used_in_historical_similarity": total_hist,
        "total_raw_signals_display_only": total_display_only,
        "total_raw_signals_missing": total_missing,
    }
    return rows, totals


def _compact_env(value: str | None) -> str:
    if not value:
        return "-"
    replacements = {
        "Mixed / Neutral": "Mixed",
        "Risk-On Rotation": "Risk-On",
        "Risk-Off": "Risk-Off",
    }
    return replacements.get(value, value[:18])


def _contributor_type(item: ScenarioContribution) -> str:
    mapping = {
        "layer_summary": "Layer",
        "raw_component": "Raw",
        "composite": "Composite",
        "regime_driver": "Regime Driver",
        "scenario_falsifier": "Falsifier",
    }
    return mapping.get(str(item.source_role or ""), str(item.source_role or "Input"))


def _short_rationale(value: str, max_len: int = 115) -> str:
    text = " ".join(str(value or "-").split())
    if len(text) <= max_len:
        return text
    return text[: max_len - 3].rstrip() + "..."


def _contributor_breakdown_table(document: DocumentObject, update: ScenarioProbabilityUpdate) -> None:
    items = sorted(update.contributions, key=lambda item: abs(item.adjusted_contribution), reverse=True)
    if not items:
        _paragraph(document, "none", compact=True)
        return
    rows = [
        [
            _contributor_type(item),
            item.input_name,
            item.parent_layer or "-",
            _signed(item.adjusted_contribution),
            _short_rationale(item.rationale),
        ]
        for item in items[:12]
    ]
    _add_table(document, ["Type", "Input", "Parent", "Contribution", "Short Rationale"], rows)


def _chart_file(chart_dir: Path, name: str) -> Path:
    chart_dir.mkdir(parents=True, exist_ok=True)
    return chart_dir / f"{name}.png"


def _plot_chart(path: Path, plotter) -> Path | None:
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None
    try:
        fig, ax = plt.subplots(figsize=(8.2, 3.4))
        plotted = plotter(ax)
        if plotted is False:
            plt.close(fig)
            return None
        fig.tight_layout()
        fig.savefig(path, dpi=150)
        plt.close(fig)
        return path
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return None


def _insert_chart(document: DocumentObject, path: Path) -> None:
    try:
        document.add_picture(str(path), width=Inches(6.4))
        _spacer(document, after=4)
    except Exception:
        _paragraph(document, f"Chart unavailable: {path.name}", compact=True)


def _scenario_probability_chart(result: MacroForecastResult, chart_dir: Path) -> Path | None:
    path = _chart_file(chart_dir, "scenario_probabilities")
    calibration_by_scenario = _calibration_by_scenario(result)

    def plot(ax):
        if result.probability_mode == "two_source_v1":
            ordered = sorted(
                result.scenario_probabilities.items(),
                key=lambda item: item[1],
                reverse=True,
            )
            if not ordered:
                return False
            labels = [_scenario_name(scenario_id)[:18] for scenario_id, _ in ordered]
            values = [probability * 100 for _, probability in ordered]
            x = list(range(len(labels)))
            ax.bar(x, values, label="Final")
            ax.set_xticks(x)
            ax.set_xticklabels(labels, rotation=35, ha="right", fontsize=8)
            ax.set_ylabel("Probability (%)")
            ax.legend(fontsize=8)
            return True
        labels = [_scenario_name(update.scenario_id)[:18] for update in result.scenario_updates]
        deterministic = [
            (calibration_by_scenario.get(update.scenario_id).deterministic_probability if calibration_by_scenario.get(update.scenario_id) else update.posterior_probability) * 100
            for update in result.scenario_updates
        ]
        historical = [
            (calibration_by_scenario.get(update.scenario_id).historical_probability if calibration_by_scenario.get(update.scenario_id) else 0.0) * 100
            for update in result.scenario_updates
        ]
        blended = [result.scenario_probabilities.get(update.scenario_id, update.posterior_probability) * 100 for update in result.scenario_updates]
        if not labels:
            return False
        x = list(range(len(labels)))
        width = 0.25
        ax.bar([i - width for i in x], deterministic, width, label="Deterministic")
        if calibration_by_scenario:
            ax.bar(x, historical, width, label="Historical")
            ax.bar([i + width for i in x], blended, width, label="Blended")
        else:
            ax.bar(x, blended, width, label="Final")
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=25, ha="right")
        ax.set_ylabel("Probability (%)")
        ax.set_title("Scenario Probabilities")
        ax.legend(fontsize=8)
        return True

    return _plot_chart(path, plot)


def _theme_support_chart(result: MacroForecastResult, chart_dir: Path) -> Path | None:
    path = _chart_file(chart_dir, "theme_macro_support")

    def plot(ax):
        themes = list(result.theme_rankings or [])
        if not themes:
            return False
        top = themes[:6]
        bottom = sorted(themes, key=lambda item: item.macro_support_score)[:3]
        selected = top + [theme for theme in bottom if theme not in top]
        labels = [theme.label[:22] for theme in selected]
        values = [theme.macro_support_score for theme in selected]
        colors = ["#2E75B6" if value >= 0 else "#C00000" for value in values]
        ax.barh(range(len(labels)), values, color=colors)
        ax.set_yticks(range(len(labels)))
        ax.set_yticklabels(labels)
        ax.axvline(0, color="#808080", linewidth=0.8)
        ax.set_xlabel("Macro support score")
        ax.set_title("Theme Macro Support: Top and Bottom")
        return True

    return _plot_chart(path, plot)


def _macro_forward_return_chart(result: MacroForecastResult, chart_dir: Path) -> Path | None:
    calibration = result.historical_calibration
    if calibration is None:
        return None
    stats_by_horizon = calibration.macro_forward_return_stats or calibration.forward_return_stats
    path = _chart_file(chart_dir, "historical_macro_forward_returns")

    def plot(ax):
        horizons = ["21d", "63d", "126d", "252d"]
        stats = [stats_by_horizon.get(horizon) for horizon in horizons]
        medians = [getattr(stat, "median", None) for stat in stats]
        if all(value is None for value in medians):
            return False
        p10 = [getattr(stat, "p10", None) if stat is not None else None for stat in stats]
        p90 = [getattr(stat, "p90", None) if stat is not None else None for stat in stats]
        x = list(range(len(horizons)))
        values = [float(value or 0.0) for value in medians]
        ax.bar(x, values, color="#2E75B6", label="Median")
        if any(v is not None for v in p10) and any(v is not None for v in p90):
            lower = [max(0.0, values[i] - float(p10[i] if p10[i] is not None else values[i])) for i in x]
            upper = [max(0.0, float(p90[i] if p90[i] is not None else values[i]) - values[i]) for i in x]
            ax.errorbar(x, values, yerr=[lower, upper], fmt="none", ecolor="#595959", capsize=4, label="P10-P90")
        ax.axhline(0, color="#808080", linewidth=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels([_horizon_label(h) for h in horizons])
        ax.set_ylabel("Forward return (%)")
        ax.set_title("Historical Macro Forward Return Profile")
        ax.legend(fontsize=8)
        return True

    return _plot_chart(path, plot)


def _risk_profile_chart(result: MacroForecastResult, chart_dir: Path) -> Path | None:
    calibration = result.historical_calibration
    if calibration is None or not calibration.risk_profile:
        return None
    risk = calibration.risk_profile
    path = _chart_file(chart_dir, "historical_risk_profile")

    def plot(ax):
        horizons = ["21d", "63d", "126d", "252d"]
        labels = [_horizon_label(h) for h in horizons]
        median_up = [risk.get(f"median_up_{h}") for h in horizons]
        median_down = [risk.get(f"median_down_{h}") for h in horizons]
        expected = [risk.get(f"expected_value_{h}") for h in horizons]
        if all(value is None for value in median_up + median_down + expected):
            return False
        x = list(range(len(horizons)))
        width = 0.25
        ax.bar([i - width for i in x], [float(v or 0.0) for v in median_up], width, label="Median Up")
        ax.bar(x, [float(v or 0.0) for v in median_down], width, label="Median Down")
        ax.bar([i + width for i in x], [float(v or 0.0) for v in expected], width, label="Expected Value")
        ax.axhline(0, color="#808080", linewidth=0.8)
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_ylabel("Return (%)")
        ax.set_title("Historical Risk Profile by Horizon")
        ax.legend(fontsize=8)
        return True

    return _plot_chart(path, plot)


def _scenario_calibration_chart(result: MacroForecastResult, chart_dir: Path) -> Path | None:
    calibration = result.historical_calibration
    if calibration is None or not calibration.scenario_calibrations:
        return None
    path = _chart_file(chart_dir, "scenario_calibration")

    def plot(ax):
        items = calibration.scenario_calibrations
        labels = [_scenario_name(item.scenario_id)[:18] for item in items]
        x = list(range(len(items)))
        width = 0.25
        ax.bar([i - width for i in x], [item.deterministic_probability * 100 for item in items], width, label="Deterministic")
        ax.bar(x, [item.historical_probability * 100 for item in items], width, label="Historical")
        ax.bar([i + width for i in x], [item.blended_probability * 100 for item in items], width, label="Blended")
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=25, ha="right")
        ax.set_ylabel("Probability (%)")
        ax.set_title("Scenario Calibration")
        ax.legend(fontsize=8)
        return True

    return _plot_chart(path, plot)


def _add_table(document: DocumentObject, headers: Sequence[str], rows: Sequence[Sequence[object]]) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        _shade_cell(header_cells[index], HEADER_FILL)
        _set_cell_text(header_cells[index], header, bold=True, size=SMALL_SIZE, color=WHITE)
    _repeat_header(table.rows[0])

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if row_index % 2:
                _shade_cell(cells[index], ALT_ROW_FILL)
            _set_cell_text(cells[index], _text(value), size=SMALL_SIZE)
    preset_widths = _table_widths_for_headers(headers)
    if preset_widths:
        _set_table_width(table, preset_widths)
    _spacer(document, after=6)
    return table


def _kv_table(document: DocumentObject, rows: Sequence[tuple[str, object]]) -> Table:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = True
    for key, value in rows:
        cells = table.add_row().cells
        _shade_cell(cells[0], KEY_FILL)
        _set_cell_text(cells[0], key, bold=True, size=SMALL_SIZE, color=PRIMARY_BLUE)
        _set_cell_text(cells[1], _text(value), size=SMALL_SIZE)
    _set_table_width(table, [1.5, 8.0])
    _spacer(document, after=6)
    return table


def _contributor_table(
    document: DocumentObject,
    contributions: Sequence[ProbabilityContribution | ScenarioContribution],
) -> None:
    if not contributions:
        _paragraph(document, "none", compact=True)
        return
    rows = [
        [
            getattr(item, "input_name", None) or getattr(item, "name", "-"),
            _contribution_source_label(item),
            _signed(
                getattr(item, "adjusted_contribution", None)
                if getattr(item, "adjusted_contribution", None) is not None
                else getattr(item, "contribution", None)
            ),
            item.rationale,
        ]
        for item in contributions
    ]
    _add_table(document, ["Input", "Source", "Contribution", "Rationale"], rows)


def _set_cell_text(
    cell: _Cell,
    value: str,
    *,
    bold: bool = False,
    size=BODY_SIZE,
    color: RGBColor | None = None,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for existing_run in list(paragraph.runs):
        paragraph._p.remove(existing_run._r)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = size
    run.font.color.rgb = color or _infer_text_color(value, bold=bold)


def _infer_text_color(value: str, *, bold: bool = False) -> RGBColor:
    stripped = value.strip()
    upper = stripped.upper()
    if stripped == "-" or upper in {"N/A", "NONE", ""}:
        return BODY_GRAY
    if stripped.startswith("+") or upper in {"BULLISH", "IMPROVING", "YES"}:
        return POSITIVE_GREEN
    if stripped.startswith("-") or upper in {"BEARISH", "DETERIORATING", "NO"}:
        return NEGATIVE_RED
    if bold:
        return RGBColor(0x40, 0x40, 0x40)
    return BODY_GRAY


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_width(table: Table, widths: Sequence[float]) -> None:
    if len(widths) != len(table.columns):
        return
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def _table_widths_for_headers(headers: Sequence[str]) -> list[float] | None:
    key = tuple(headers)
    presets: dict[tuple[str, ...], list[float]] = {
        (
            "Group",
            "Avg Similarity",
            "Features Used",
            "Features Missing",
            "Coverage",
            "Used Feature IDs",
            "Missing Feature IDs",
            "Interpretation",
        ): [1.00, 0.60, 0.55, 0.65, 0.55, 1.75, 1.75, 1.90],
        (
            "Scenario",
            "Deterministic Probability",
            "Historical Probability",
            "Blended Probability",
            "Analog Effect",
            "Supporting Analogues",
            "Confidence",
            "Rationale",
        ): [1.35, 0.75, 0.75, 0.75, 0.65, 0.65, 0.55, 4.20],
        (
            "Date",
            "Weight",
            "Env",
            "Score",
            "VIX",
            "Breadth",
            "63D Fwd",
            "21D Fwd",
            "Scenario",
            "Map Tag",
            "V1",
            "V2",
            "Blend",
            "Strong Groups",
            "Weak Groups",
        ): [0.65, 0.45, 0.55, 0.45, 0.45, 0.45, 0.50, 0.50, 1.05, 0.85, 0.45, 0.45, 0.45, 1.20, 1.20],
        (
            "Horizon",
            "N",
            "Weight Sum",
            "Median",
            "Mean",
            "% Positive",
            "P10",
            "P25",
            "P75",
            "P90",
            "Worst",
            "Best",
        ): [0.80, 0.35, 0.55, 0.50, 0.50, 0.60, 0.45, 0.45, 0.45, 0.45, 0.50, 0.50],
        (
            "Horizon",
            "Win Rate",
            "Median Up",
            "Median Down",
            "Expected Value",
            "P10 Return",
            "P90 Return",
            "Worst Return",
        ): [0.85, 0.65, 0.70, 0.75, 0.85, 0.70, 0.70, 0.80],
        (
            "Input",
            "Value",
            "Source",
            "Observed Date",
            "Lookback",
            "Freshness",
            "Audit Status",
        ): [1.30, 0.65, 2.05, 0.70, 1.20, 0.75, 2.90],
        (
            "Layer",
            "Score",
            "Status",
            "Confidence/Data Quality",
            "Used in Math?",
            "Key Signals",
            "Raw Components Attached",
        ): [0.70, 0.55, 0.80, 0.85, 0.55, 4.35, 1.90],
        (
            "Type",
            "Input",
            "Parent",
            "Contribution",
            "Short Rationale",
        ): [0.85, 1.45, 0.80, 0.70, 5.70],
        (
            "Scenario",
            "Probability",
            "Exposure",
            "Contribution",
            "Rationale",
        ): [1.45, 0.75, 0.70, 0.75, 5.75],
        (
            "Rank",
            "Ticker",
            "Score",
            "Top Drivers",
            "Formula Notes",
        ): [0.35, 0.75, 0.55, 3.55, 4.50],
        (
            "Rank",
            "Factor",
            "Score",
            "Top Drivers",
            "Formula Notes",
        ): [0.35, 1.05, 0.55, 3.35, 4.20],
    }
    return presets.get(key)


def _repeat_header(row: _Row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _section(document: DocumentObject, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = SECTION_SIZE
    run.font.color.rgb = PRIMARY_BLUE


def _subsection(document: DocumentObject, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = SUBSECTION_SIZE
    run.font.color.rgb = SECONDARY_BLUE


def _paragraph(document: DocumentObject, text: str, *, bold: bool = False, compact: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2 if compact else 5)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = PRIMARY_BLUE if bold else BODY_GRAY


def _spacer(document: DocumentObject, *, after: int = 6) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)


def _scenario_name(scenario_id: str) -> str:
    labels = {
        "reopening_soft_landing": "Reopening / Soft Landing",
        "sticky_late_cycle_ai": "Sticky Late Cycle AI",
        "oil_inflation_tail": "Oil Inflation Tail",
        "late_cycle_risk_off": "Late Cycle Risk-Off",
        "ai_capex_rollover": "AI Capex Rollover",
    }
    return labels.get(scenario_id, scenario_id.replace("_", " ").title())


def _to_float(value: object) -> float | None:
    if value is None or isinstance(value, bool):
        return None
    try:
        output = float(value)
    except (TypeError, ValueError):
        return None
    if output != output:
        return None
    return output


def _pct(value: float | None) -> str:
    numeric = _to_float(value)
    if numeric is None:
        return "n/a"
    return f"{numeric:.2%}"


def _decimal(value: float | None, *, places: int = 2) -> str:
    numeric = _to_float(value)
    if numeric is None:
        return "n/a"
    return f"{numeric:.{places}f}"


def _signed(value: float | None) -> str:
    numeric = _to_float(value)
    if numeric is None:
        return "n/a"
    return f"{numeric:+.2f}"


def _text(value: object) -> str:
    if value is None:
        return "n/a"
    date_text = _format_date_for_display(value)
    if date_text is not None:
        return date_text
    return str(value)


def _format_date_for_display(value: object) -> str | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.strftime("%m/%d/%y")
    if isinstance(value, date):
        return value.strftime("%m/%d/%y")
    if not isinstance(value, str):
        return None
    text = value.strip()
    try:
        if ISO_DATE_RE.match(text):
            return datetime.strptime(text, "%Y-%m-%d").strftime("%m/%d/%y")
        if ISO_DATETIME_RE.match(text):
            return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%m/%d/%y")
    except ValueError:
        return None
    return None


def _yes_no(value: bool) -> str:
    return "yes" if value else "no"


def _contribution_source_label(item: ProbabilityContribution | ScenarioContribution) -> str:
    role = getattr(item, "source_role", None) or "input"
    layer = getattr(item, "parent_layer", None)
    label = f"{role}/{layer}" if layer else role
    if getattr(item, "capped_by_dedupe", False):
        label += ", capped"
    return label


def _top_scenario_effects(signal: MacroInputSignal) -> str:
    if not signal.affected_scenarios:
        return "none"
    return ", ".join(
        f"{impact.scenario_id} {impact.direction} {impact.strength:.2f}"
        for impact in signal.affected_scenarios[:3]
    )


def _raw_components_for_layer(result: MacroForecastResult, parent_layer: str | None) -> str:
    if result.forecast_input_set is None or parent_layer is None:
        return "-"
    values = [
        signal.input_id
        for signal in result.forecast_input_set.raw_component_signals
        if signal.parent_layer == parent_layer
    ]
    return ", ".join(values) if values else "-"


def _join(items: Iterable[object]) -> str:
    values = [_text(item) for item in items if item is not None and str(item)]
    return ", ".join(values) if values else "-"


def _horizon_count_text(values: object) -> str:
    if not isinstance(values, dict):
        return "-"
    ordered = ["1d", "5d", "10d", "21d", "63d", "126d", "252d"]
    parts = [
        f"{horizon}: {values.get(horizon, 0)}"
        for horizon in ordered
        if values.get(horizon, 0)
    ]
    return "; ".join(parts) if parts else "none"


def _bullet_text(items: Iterable[object]) -> str:
    values = [_text(item) for item in items if item is not None and str(item)]
    return "\n".join(f"- {item}" for item in values) if values else "-"


def _signal_reading(value: float | str | bool | None) -> str:
    if value is None:
        return "n/a"
    if isinstance(value, float):
        return f"{value:.2f}"
    return _text(value)


def _floor_cap_label(update: ScenarioProbabilityUpdate) -> str:
    labels: list[str] = []
    if update.floor_applied:
        labels.append(f"floor {_pct(update.floor_value)}")
    if update.cap_applied:
        labels.append(f"cap {_pct(update.cap_value)}")
    return ", ".join(labels) if labels else "-"


def _calibration_by_scenario(result: MacroForecastResult):
    calibration = result.historical_calibration
    if calibration is None:
        return {}
    return {
        item.scenario_id: item
        for item in calibration.scenario_calibrations
    }


def _top_driver(update: ScenarioProbabilityUpdate) -> str:
    contributors = list(update.top_positive_contributors) + list(update.top_negative_contributors)
    if not contributors:
        return "none"
    top = sorted(contributors, key=lambda item: abs(item.contribution), reverse=True)[0]
    return f"{top.name} ({_signed(top.contribution)})"


def _top_contributions(
    update: ScenarioProbabilityUpdate,
    *,
    positive: bool,
) -> list[ScenarioContribution]:
    if positive:
        items = [item for item in update.contributions if item.adjusted_contribution > 0]
    else:
        items = [item for item in update.contributions if item.adjusted_contribution < 0]
    return sorted(items, key=lambda item: abs(item.adjusted_contribution), reverse=True)[:5]


def _contributions_for_roles(
    update: ScenarioProbabilityUpdate,
    roles: set[str],
) -> list[ScenarioContribution]:
    return sorted(
        [
            item
            for item in update.contributions
            if str(item.source_role or "") in roles
        ],
        key=lambda item: abs(item.adjusted_contribution),
        reverse=True,
    )[:5]


def _ranking_driver_summary(contributions: Sequence[RankingContribution]) -> str:
    if not contributions:
        return "none"
    return ", ".join(
        f"{item.source_label} {item.contribution:+.2f}"
        for item in contributions[:3]
    )


def _theme_contribution_summary(theme) -> str:
    contributions = getattr(theme, "scenario_contributions", [])
    if not contributions:
        return "none"
    return "; ".join(
        f"{item.scenario_label} {item.contribution:+.2f}"
        for item in sorted(
            contributions,
            key=lambda value: abs(value.contribution),
            reverse=True,
        )
    )


def _evidence_text(evidence_items: Iterable[object]) -> str:
    lines: list[str] = []
    for evidence in evidence_items:
        claim = getattr(evidence, "claim", "")
        computation = getattr(evidence, "computation", "")
        upstream_claims = getattr(evidence, "upstream_claims", [])
        parts = [part for part in [claim, computation, _join(upstream_claims)] if part and part != "-"]
        if parts:
            lines.append(" | ".join(parts))
    return _bullet_text(lines)
