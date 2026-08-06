from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

from src.agent_system.diagnostics import input_ingestion_audit
from src.agent_system.diagnostics.input_ingestion_audit import (
    _audit_statuses,
    build_input_ingestion_audit,
    calculate_freshness,
    format_auditable_layer_key_signal,
    input_audit_warnings_from_input_set,
    provenance_summary_rows_from_input_set,
)
from src.agent_system.forecasting import macro_forecast_runner as runner
from src.agent_system.forecasting.input_signals import build_forecast_input_set
from src.agent_system.forecasting.macro_forecast_runner import run_macro_forecast
from src.agent_system.orchestration.stub_agents import make_stub_regime_state
from src.agent_system.reporting.macro_forecast_docx import generate_macro_forecast_docx
from src.agent_system.schemas.macro_forecast import MacroInputSignal
from src.state.market_state import MarketState
from src.state.regime_data import RegimeInputs

from two_source_fixtures import patch_two_source_runner


def _raw_inputs() -> RegimeInputs:
    return RegimeInputs(
        asof_date="2026-06-05",
        net_liquidity_z=0.7,
        nfci=-0.51,
        nfci_inverted=0.82,
        m2_growth_yoy=4.72,
        hy_spread_level=280,
        hy_spread_z=-0.64,
        hy_spread_chg_4w=0,
        ig_spread_level=75,
        vix_level=22.41,
        vix_z_20d=2.11,
        vix_term_slope=0.81,
        put_call_ratio=1.11,
        put_call_5d_ma=1.11,
        pct_above_200d=55,
        sectors_green=7,
        rsp_vs_spy_z=-0.13,
    )


def _market_state() -> MarketState:
    return MarketState(
        asof_utc="2026-06-05T00:00:00+00:00",
        horizon="3M",
        cross_asset_returns={
            "SPY": 9.56,
            "RSP": 7.78,
            "QQQ": 11.0,
            "IWM": 6.2,
            "HYG": 2.1,
            "TLT": -0.4,
            "GLD": 1.2,
            "USO": 3.0,
            "BTC-USD": 4.0,
        },
        sector_returns={"XLK": 1.0, "XLE": 0.5},
        leadership_top3=[("Technology", 1.0), ("Energy", 0.5), ("Industrials", 0.2)],
        sectors_green=7,
        dispersion=0.4,
        spy_clv=0.6,
        spy_range_pct=0.8,
        market_session_date="2026-06-05",
        spy_vol_z_20d=1.1,
        volume_confirmation=0.7,
        spy_above_vwap=True,
        spy_above_prev_close=True,
        vix_level=22.41,
        vix_z_20d=2.11,
        vix_change_pct_1d=-3.0,
    )


def _history_df() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "date": pd.Timestamp("2026-06-03"),
                "vix_level": 22.0,
                "hy_spread_level": 285.0,
                "rsp_minus_spy": -1.0,
                "put_call_ratio": 1.0,
            },
            {
                "date": pd.Timestamp("2026-06-04"),
                "vix_level": 22.5,
                "hy_spread_level": 280.0,
                "rsp_minus_spy": -1.2,
                "put_call_ratio": 1.1,
            },
        ]
    )


def _forecast_input_set():
    return build_forecast_input_set(
        make_stub_regime_state(),
        raw_inputs=_raw_inputs(),
        market_state=_market_state(),
        horizon="3m",
    )


def _patch_audit_loaders(monkeypatch) -> None:
    monkeypatch.setattr(input_ingestion_audit, "_load_regime_for_cli", lambda asof_date=None: make_stub_regime_state())
    monkeypatch.setattr(input_ingestion_audit, "_load_regime_inputs_for_cli", lambda asof_date=None: _raw_inputs())
    monkeypatch.setattr(input_ingestion_audit, "_load_market_state_for_cli", lambda asof_date=None, horizon="3m": _market_state())
    monkeypatch.setattr(input_ingestion_audit, "_load_historical_df", lambda: _history_df())


def _document_text(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_key_signal_includes_value_lookback_and_source():
    input_set = _forecast_input_set()
    breadth = next(signal for signal in input_set.layer_summary_signals if signal.parent_layer == "breadth")
    positioning = next(signal for signal in input_set.layer_summary_signals if signal.parent_layer == "positioning")

    breadth_key = format_auditable_layer_key_signal(input_set, breadth, horizon="3m")
    positioning_key = format_auditable_layer_key_signal(input_set, positioning, horizon="3m")

    assert "RSP-SPY: -1.78 pp over 3M / 63 trading days" in breadth_key
    assert "RSP +7.78%" in breadth_key
    assert "SPY +9.56%" in breadth_key
    assert "source=MarketState.cross_asset_returns" in breadth_key
    assert "Generic put/call ratio: 1.11" in positioning_key
    assert "Cboe equity unresolved" in positioning_key
    assert "observed_date=" in positioning_key


def test_put_call_disambiguation_and_duplication_warnings():
    input_set = _forecast_input_set()

    rows = provenance_summary_rows_from_input_set(input_set, asof_date="2026-06-05", horizon="3m")
    put_call = next(row for row in rows if row["input_id"] == "put_call_ratio")
    warnings = input_audit_warnings_from_input_set(input_set, horizon="3m")

    assert "source_ambiguous" in put_call["audit_status"]
    assert "duplicated_value" in put_call["audit_status"]
    assert "put_call_ratio source unresolved" in put_call["warnings"]
    assert any("Cboe equity put/call not ingested" in warning for warning in warnings)
    assert any("put_call_ratio equals put_call_5D MA" in warning or "put_call_ratio equals put_call_5d_ma" in warning for warning in warnings)


def test_cboe_equity_put_call_labels_correctly_when_explicit():
    input_set = _forecast_input_set()
    equity_signal = MacroInputSignal(
        input_id="cboe_equity_put_call_ratio",
        name="Cboe equity put/call ratio",
        category="volatility",
        current_value=0.44,
        unit="ratio",
        percentile=None,
        z_score=None,
        trend="stable",
        signal="neutral",
        confidence=0.9,
        data_quality="high",
        raw_value=0.44,
        source_object="Cboe",
        parent_layer="volatility",
        role="raw_component",
    )
    input_set = input_set.model_copy(
        update={
            "raw_component_signals": [*input_set.raw_component_signals, equity_signal],
            "all_signals": [*input_set.all_signals, equity_signal],
        }
    )
    positioning = next(signal for signal in input_set.layer_summary_signals if signal.parent_layer == "positioning")

    key = format_auditable_layer_key_signal(input_set, positioning, horizon="3m")

    assert "Equity put/call: 0.44" in key
    assert "Cboe Daily Market Statistics" in key


def test_freshness_rules_mark_daily_stale_and_monthly_acceptable():
    daily_status, daily_lag, _ = calculate_freshness("2026-06-01", "2026-06-05", "daily")
    monthly_status, monthly_lag, _ = calculate_freshness("2026-05-05", "2026-06-05", "monthly")

    assert daily_status == "stale"
    assert daily_lag == 4.0
    assert monthly_status in {"fresh", "acceptable_lag"}
    assert monthly_lag == 31.0


def test_derived_signal_without_lookback_is_flagged():
    status, warnings = _audit_statuses(
        input_id="rsp_minus_spy",
        value=-1.2,
        source_object="MarketState",
        source_field="cross_asset_returns.RSP-SPY",
        observed_date="2026-06-05",
        freshness_status="fresh",
        lookback_window=None,
        calculation_method="RSP minus SPY",
        duplicate_put_call=False,
        has_explicit_cboe_equity=False,
    )

    assert "missing_lookback" in status
    assert warnings == []


def test_ingestion_audit_saves_csv_xlsx_json_and_required_columns(monkeypatch, tmp_path):
    _patch_audit_loaders(monkeypatch)

    matrix = build_input_ingestion_audit(output_dir=str(tmp_path), save_csv=True, save_xlsx=True, save_json=True)

    for column in [
        "input_id",
        "display_label",
        "provider",
        "source_object",
        "source_field",
        "observed_date",
        "lookback_window",
        "calculation_method",
        "freshness_status",
        "audit_status",
    ]:
        assert column in matrix.columns
    paths = matrix.attrs["output_paths"]
    assert Path(paths["csv"]).exists()
    assert Path(paths["xlsx"]).exists()
    assert Path(paths["json"]).exists()
    payload = json.loads(Path(paths["json"]).read_text(encoding="utf-8"))
    assert "inputs" in payload
    with zipfile.ZipFile(paths["xlsx"]) as zf:
        workbook_xml = zf.read("xl/workbook.xml").decode("utf-8")
    assert "All Inputs" in workbook_xml
    assert "Warnings" in workbook_xml
    assert "Group Summary" in workbook_xml


def test_report_contains_input_provenance_and_auditable_key_signals(monkeypatch, tmp_path):
    patch_two_source_runner(monkeypatch, runner)
    result = run_macro_forecast(
        make_stub_regime_state(),
        raw_inputs=_raw_inputs(),
        market_state=_market_state(),
    )

    path = generate_macro_forecast_docx(result, tmp_path / "macro_forecast.docx")
    text = _document_text(path)

    assert "Input Provenance Summary" in text
    assert "Input Audit Warnings" in text
    assert "Full input ingestion audit saved to:" in text
    assert "RSP-SPY: -1.78 pp over 3M / 63 trading days" in text
    assert "Generic put/call ratio: 1.11" in text
    assert "Cboe equity unresolved" in text
