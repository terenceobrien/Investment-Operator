from __future__ import annotations

from pathlib import Path

from docx import Document

from src.agent_system.forecasting import macro_forecast_runner as runner
from src.agent_system.forecasting.macro_forecast_runner import run_macro_forecast
from src.agent_system.orchestration.stub_agents import make_stub_regime_state
from src.agent_system.reporting.macro_forecast_docx import generate_macro_forecast_docx
from src.agent_system.schemas.macro_forecast import MacroForecastResult

from two_source_fixtures import analogue_evidence_fixture, patch_two_source_runner


def _document_text(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_macro_forecast_result_json_round_trips_under_two_source(monkeypatch, tmp_path):
    patch_two_source_runner(monkeypatch, runner)
    result = run_macro_forecast(make_stub_regime_state())
    fixture_path = tmp_path / "macro_forecast_result.json"
    fixture_path.write_text(result.model_dump_json(), encoding="utf-8")

    parsed = MacroForecastResult.model_validate_json(fixture_path.read_text(encoding="utf-8"))

    assert parsed.probability_mode == "two_source_v1"
    assert parsed.scenario_updates == []
    assert parsed.mixture_report["combination"] == "linear_mixture"


def test_docx_generation_includes_two_source_mixture_sections(monkeypatch, tmp_path):
    patch_two_source_runner(
        monkeypatch,
        runner,
        evidence=analogue_evidence_fixture(trailing_max=0.41, stress_advisory=True),
    )
    result = run_macro_forecast(make_stub_regime_state())

    path = generate_macro_forecast_docx(result, tmp_path / "macro_forecast.docx")
    text = _document_text(path)

    assert path.exists()
    assert "Scenario Probabilities" in text
    assert "BVAR Soft" in text
    assert "Analogue Implied" in text
    assert "Scenario Probability Math" in text
    assert "BVAR Model Limitations" in text
    assert "Monitoring — no probability impact" in text
    assert "Legacy rolling historical calibration is retired" in text
